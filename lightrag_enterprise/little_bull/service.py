from __future__ import annotations

import asyncio
import copy
import hashlib
import json
import os
import re
import shutil
from datetime import date, datetime, timezone
from functools import partial
from io import BytesIO
from pathlib import Path
from typing import Any

import aiofiles
from fastapi import BackgroundTasks, HTTPException, Response, UploadFile, status

from lightrag.base import QueryParam, StoragesStatus
from lightrag.llm.openai import openai_complete_if_cache
from lightrag.namespace import NameSpace
from lightrag.utils import generate_track_id

from lightrag_enterprise.system import (
    ACTIVITY_ACTIVITY_READ,
    ACTIVITY_AGENT_MANAGE,
    ACTIVITY_AREA_READ,
    ACTIVITY_ASSISTANTS_READ,
    ACTIVITY_DOCUMENT_DELETE,
    ACTIVITY_DOCUMENT_READ,
    ACTIVITY_DOCUMENT_REINDEX,
    ACTIVITY_DOCUMENT_UPLOAD,
    ACTIVITY_CONVERSATION_EXPORT,
    ACTIVITY_CONVERSATION_READ,
    ACTIVITY_CONVERSATION_SAVE,
    ACTIVITY_CORRELATION_DECIDE,
    ACTIVITY_CORRELATION_SUGGEST,
    ACTIVITY_MODEL_MANAGE,
    ACTIVITY_QUERY,
    ACTIVITY_WORKSPACE_MANAGE,
    AccessControlService,
    ApprovalService,
    ApprovalStatus,
    AuditService,
    Principal,
    Workspace,
)
from lightrag_enterprise.system.policy_keys import (
    PRIVATE_DATA_HOSTED_LLM_EXCEPTION_POLICY,
    WORKSPACE_DATA_PLANE_POLICY,
    WORKSPACE_PRIVATE_POLICY,
)
from lightrag_enterprise.system.runtime import approvals_enforced, private_strict_enabled
from lightrag_enterprise.system.models import utc_now

from .models import (
    AgentBuilderSession,
    AgentContextBudget,
    Backlink,
    CanvasBoard,
    CanvasEdge,
    CanvasNode,
    ContentMap,
    KnowledgeDossier,
    KnowledgeInboxItem,
    KnowledgeTrail,
    KnowledgeTrailStep,
    DailyNote,
    LittleBullActivityItem,
    LittleBullAgentBuilderPublishRequest,
    LittleBullAgentBuilderSessionRequest,
    LittleBullAgentConfig,
    LittleBullAgentContextBudgetRequest,
    LittleBullAgentStudioPreviewRequest,
    LittleBullAgentStudioPreviewResponse,
    LittleBullArea,
    LittleBullAssistant,
    LittleBullBacklinkRequest,
    LittleBullCanvasAnalysis,
    LittleBullCanvasBoardDetail,
    LittleBullCanvasBoardRequest,
    LittleBullCanvasEdgeRequest,
    LittleBullCanvasNodeRequest,
    LittleBullContentMapRequest,
    LittleBullConversation,
    LittleBullConversationSaveRequest,
    LittleBullCorrelationSuggestion,
    LittleBullCorrelationSuggestionRequest,
    LittleBullDocument,
    LittleBullDocumentsResponse,
    LittleBullEmbeddingCatalogItem,
    LittleBullEmbeddingCostEstimateRequest,
    LittleBullEmbeddingCostEstimateResponse,
    LittleBullKnowledgeBase,
    LittleBullKnowledgeBaseAttachResponse,
    LittleBullKnowledgeBaseReindexRequest,
    LittleBullKnowledgeBaseReindexResponse,
    LittleBullKnowledgeBaseRollbackRequest,
    LittleBullKnowledgeBaseRollbackResponse,
    LittleBullKnowledgeBaseUpsertRequest,
    LittleBullKnowledgeGroupRequest,
    LittleBullDailyNoteRequest,
    LittleBullInboxItemRequest,
    LittleBullInboxItemStatusRequest,
    LittleBullKnowledgeSubgroupRequest,
    LittleBullKnowledgeTrailDetail,
    LittleBullKnowledgeTrailRequest,
    LittleBullKnowledgeTrailStepRequest,
    LittleBullMarkdownNoteRequest,
    LittleBullMarkdownNoteResponse,
    LittleBullModelSetting,
    LittleBullProvenancePanel,
    LittleBullQueryRequest,
    LittleBullQueryResponse,
    LittleBullReindexArchivedResponse,
    LittleBullSourceProvenanceRequest,
    LittleBullUploadResponse,
    KnowledgeGroup,
    KnowledgeSubgroup,
    MarkdownNote,
    NoteRegistry,
    SourceProvenance,
    TagRegistry,
    WikiLink,
)
from .admin_store import LittleBullAdminStore
from .agent_studio import (
    agent_studio_preview,
    build_agent_studio_prompt,
    normalize_agent_studio_config,
    normalize_tool_id,
    validate_agent_studio_config,
)
from .private_gateway import PrivateLocalGateway
from .model_catalog import (
    OPENROUTER_EMBEDDING_HOST,
    embedding_catalog,
    estimate_embedding_cost,
    estimate_tokens_from_characters,
    estimate_tokens_from_pages,
    find_embedding_model,
)


AREA_PRESENTATION = {
    "default": ("Default", "📁", "#2563EB"),
    "casa": ("Casa", "🏠", "#FACC15"),
    "familia": ("Família", "👨‍👩‍👧", "#F97316"),
    "financas": ("Finanças", "💳", "#22C55E"),
    "trabalho": ("Trabalho", "💼", "#2563EB"),
    "estudos": ("Estudos", "📚", "#7C3AED"),
    "negocio": ("Pequeno negócio", "🧾", "#0F766E"),
}

WIKI_LINK_RE = re.compile(r"\[\[([^\]\n|]+)(?:\|([^\]\n]+))?\]\]")
MARKDOWN_TAG_RE = re.compile(r"(?<![\w/])#([\w][\w-]{0,63})")


def slugify_workspace(value: str) -> str:
    normalized = re.sub(r"[^a-zA-Z0-9_]+", "_", value.strip().lower())
    normalized = re.sub(r"_+", "_", normalized).strip("_")
    return normalized or "base"


def extract_wiki_links(markdown: str) -> list[dict[str, str]]:
    links: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for match in WIKI_LINK_RE.finditer(markdown):
        target_label = match.group(1).strip()
        link_text = (match.group(2) or target_label).strip()
        if not target_label:
            continue
        key = (target_label.casefold(), link_text.casefold())
        if key in seen:
            continue
        seen.add(key)
        links.append({"target_label": target_label, "link_text": link_text})
    return links


def extract_markdown_tags(markdown: str) -> list[str]:
    tags: list[str] = []
    seen: set[str] = set()
    for match in MARKDOWN_TAG_RE.finditer(markdown):
        tag = f"#{match.group(1).strip().lower()}"
        if tag in seen:
            continue
        seen.add(tag)
        tags.append(tag)
    return tags


def markdown_summary(markdown: str, limit: int = 240) -> str:
    lines = [line.strip() for line in markdown.splitlines() if line.strip()]
    summary = " ".join(lines)
    summary = WIKI_LINK_RE.sub(lambda match: (match.group(2) or match.group(1)).strip(), summary)
    summary = re.sub(r"\s+", " ", summary).strip(" #")
    return summary[:limit]


def canonical_ref_kind(ref_kind: str) -> str:
    normalized = ref_kind.strip().lower()
    if normalized in {"doc", "document"}:
        return "document"
    if normalized in {"note", "markdown_note"}:
        return "note"
    return normalized


def sanitize_upload_filename(filename: str, input_dir: Path) -> str:
    clean_name = filename.replace("/", "").replace("\\", "").replace("..", "")
    clean_name = "".join(char for char in clean_name if ord(char) >= 32 and char != "\x7f")
    clean_name = clean_name.strip().strip(".")
    if not clean_name:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid filename")
    try:
        final_path = (input_dir / clean_name).resolve()
        if not final_path.is_relative_to(input_dir.resolve()):
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsafe filename")
    except OSError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid filename") from exc
    return clean_name


def unique_input_filename(
    filename: str,
    input_dir: Path,
    *,
    reserved_names: set[str] | None = None,
    extra_dirs: list[Path] | None = None,
) -> str:
    safe_name = sanitize_upload_filename(filename, input_dir)
    reserved_names = reserved_names or set()
    extra_dirs = extra_dirs or []

    def is_available(name: str) -> bool:
        if name in reserved_names:
            return False
        if (input_dir / name).exists():
            return False
        return not any((directory / name).exists() for directory in extra_dirs)

    if is_available(safe_name):
        return safe_name

    stem = Path(safe_name).stem or "document"
    suffix = Path(safe_name).suffix
    for index in range(1, 1000):
        next_name = f"{stem}_reindex_{index:03d}{suffix}"
        if is_available(next_name):
            return next_name
    raise HTTPException(
        status_code=status.HTTP_409_CONFLICT,
        detail="Could not create a unique recovery filename.",
    )


def coerce_payload_bool(value: Any, default: bool) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "on"}
    return bool(value)


class LittleBullService:
    def __init__(
        self,
        *,
        rag: Any,
        doc_manager: Any,
        repository: Any,
        access: AccessControlService,
        audit: AuditService,
        approvals: ApprovalService,
        private_gateway: PrivateLocalGateway | None = None,
        admin_store: LittleBullAdminStore | None = None,
    ) -> None:
        self.rag = rag
        self.doc_manager = doc_manager
        self.repository = repository
        self.access = access
        self.audit = audit
        self.approvals = approvals
        self.private_gateway = private_gateway or PrivateLocalGateway(rag)
        self.admin_store = admin_store or LittleBullAdminStore()

    def _require(self, principal: Principal, activity: str, workspace_id: str | None = None) -> None:
        decision = self.access.require(principal, activity=activity, workspace_id=workspace_id)
        if not decision.allowed:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=decision.reason)

    def _require_master(self, principal: Principal) -> None:
        if not principal.is_master_global:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="MASTER global required.")

    def _require_backed_workspace(self, workspace_id: str) -> None:
        current_workspace = getattr(self.rag, "workspace", None) or "default"
        if workspace_id != current_workspace:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=(
                    f"Workspace '{workspace_id}' is authorized but is not attached to the current "
                    f"LightRAG data plane '{current_workspace}'."
                ),
            )

    def _workspace_rag_cache(self) -> dict[str, Any]:
        cache = getattr(self.rag, "_little_bull_workspace_rags", None)
        if cache is None:
            cache = {}
            setattr(self.rag, "_little_bull_workspace_rags", cache)
        return cache

    async def _data_plane_policy(self, workspace_id: str) -> dict[str, Any] | None:
        tenant_id = await self._workspace_tenant(workspace_id)
        policy = await self.repository.get_policy(
            WORKSPACE_DATA_PLANE_POLICY,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        return policy if isinstance(policy, dict) else None

    async def _data_plane_attached(self, workspace_id: str) -> bool:
        if self._is_backed_workspace(workspace_id):
            return True
        policy = await self._data_plane_policy(workspace_id)
        return bool(policy and policy.get("attached"))

    async def _require_data_plane(self, workspace_id: str) -> Any:
        if workspace_id == self._current_workspace_id():
            return self.rag
        cache = self._workspace_rag_cache()
        if workspace_id in cache:
            return cache[workspace_id]
        if not await self._data_plane_attached(workspace_id):
            current_workspace = self._current_workspace_id()
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=(
                    f"Workspace '{workspace_id}' is authorized but is not attached to a LightRAG data plane. "
                    f"Current runtime workspace is '{current_workspace}'."
                ),
            )
        return await self._ensure_workspace_data_plane(workspace_id)

    async def _ensure_workspace_data_plane(self, workspace_id: str) -> Any:
        if workspace_id == self._current_workspace_id():
            return self.rag
        cache = self._workspace_rag_cache()
        if workspace_id in cache:
            return cache[workspace_id]
        workspace = await self.repository.get_workspace(workspace_id)
        if workspace is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f"Workspace '{workspace_id}' not found.")
        try:
            workspace_rag = await self._create_workspace_rag(workspace_id)
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=f"Could not attach LightRAG data plane for workspace '{workspace_id}'.",
            ) from exc
        cache[workspace_id] = workspace_rag
        return workspace_rag

    async def _create_workspace_rag(self, workspace_id: str) -> Any:
        if hasattr(self.rag, "clone_for_workspace"):
            clone = self.rag.clone_for_workspace(workspace_id)
            if hasattr(clone, "__await__"):
                clone = await clone
            return clone

        workspace_rag = copy.copy(self.rag)
        workspace_rag.workspace = workspace_id

        if not all(
            hasattr(workspace_rag, attr)
            for attr in (
                "key_string_value_json_storage_cls",
                "vector_db_storage_cls",
                "graph_storage_cls",
                "doc_status_storage_cls",
                "initialize_storages",
            )
        ):
            return workspace_rag

        global_config = self._workspace_global_config(workspace_id)
        workspace_rag.key_string_value_json_storage_cls = self._rebind_storage_factory(
            self.rag.key_string_value_json_storage_cls,
            global_config,
        )
        workspace_rag.vector_db_storage_cls = self._rebind_storage_factory(
            self.rag.vector_db_storage_cls,
            global_config,
        )
        workspace_rag.graph_storage_cls = self._rebind_storage_factory(
            self.rag.graph_storage_cls,
            global_config,
        )
        workspace_rag.llm_response_cache = workspace_rag.key_string_value_json_storage_cls(
            namespace=NameSpace.KV_STORE_LLM_RESPONSE_CACHE,
            workspace=workspace_id,
            global_config=global_config,
            embedding_func=workspace_rag.embedding_func,
        )
        workspace_rag.text_chunks = workspace_rag.key_string_value_json_storage_cls(
            namespace=NameSpace.KV_STORE_TEXT_CHUNKS,
            workspace=workspace_id,
            embedding_func=workspace_rag.embedding_func,
        )
        workspace_rag.full_docs = workspace_rag.key_string_value_json_storage_cls(
            namespace=NameSpace.KV_STORE_FULL_DOCS,
            workspace=workspace_id,
            embedding_func=workspace_rag.embedding_func,
        )
        workspace_rag.full_entities = workspace_rag.key_string_value_json_storage_cls(
            namespace=NameSpace.KV_STORE_FULL_ENTITIES,
            workspace=workspace_id,
            embedding_func=workspace_rag.embedding_func,
        )
        workspace_rag.full_relations = workspace_rag.key_string_value_json_storage_cls(
            namespace=NameSpace.KV_STORE_FULL_RELATIONS,
            workspace=workspace_id,
            embedding_func=workspace_rag.embedding_func,
        )
        workspace_rag.entity_chunks = workspace_rag.key_string_value_json_storage_cls(
            namespace=NameSpace.KV_STORE_ENTITY_CHUNKS,
            workspace=workspace_id,
            embedding_func=workspace_rag.embedding_func,
        )
        workspace_rag.relation_chunks = workspace_rag.key_string_value_json_storage_cls(
            namespace=NameSpace.KV_STORE_RELATION_CHUNKS,
            workspace=workspace_id,
            embedding_func=workspace_rag.embedding_func,
        )
        workspace_rag.chunk_entity_relation_graph = workspace_rag.graph_storage_cls(
            namespace=NameSpace.GRAPH_STORE_CHUNK_ENTITY_RELATION,
            workspace=workspace_id,
            embedding_func=workspace_rag.embedding_func,
        )
        workspace_rag.entities_vdb = workspace_rag.vector_db_storage_cls(
            namespace=NameSpace.VECTOR_STORE_ENTITIES,
            workspace=workspace_id,
            embedding_func=workspace_rag.embedding_func,
            meta_fields={"entity_name", "source_id", "content", "file_path"},
        )
        workspace_rag.relationships_vdb = workspace_rag.vector_db_storage_cls(
            namespace=NameSpace.VECTOR_STORE_RELATIONSHIPS,
            workspace=workspace_id,
            embedding_func=workspace_rag.embedding_func,
            meta_fields={"src_id", "tgt_id", "source_id", "content", "file_path"},
        )
        workspace_rag.chunks_vdb = workspace_rag.vector_db_storage_cls(
            namespace=NameSpace.VECTOR_STORE_CHUNKS,
            workspace=workspace_id,
            embedding_func=workspace_rag.embedding_func,
            meta_fields={"full_doc_id", "content", "file_path"},
        )
        workspace_rag.doc_status = workspace_rag.doc_status_storage_cls(
            namespace=NameSpace.DOC_STATUS,
            workspace=workspace_id,
            global_config=global_config,
            embedding_func=None,
        )
        workspace_rag._storages_status = StoragesStatus.CREATED
        await workspace_rag.initialize_storages()
        return workspace_rag

    def _workspace_global_config(self, workspace_id: str) -> dict[str, Any]:
        source = {}
        for candidate in (
            getattr(getattr(self.rag, "full_docs", None), "global_config", None),
            getattr(getattr(self.rag, "doc_status", None), "global_config", None),
        ):
            if isinstance(candidate, dict):
                source = candidate
                break
        global_config = dict(source)
        global_config["workspace"] = workspace_id
        global_config.setdefault("working_dir", getattr(self.rag, "working_dir", "./rag_storage"))
        return global_config

    @staticmethod
    def _rebind_storage_factory(factory: Any, global_config: dict[str, Any]) -> Any:
        if isinstance(factory, partial):
            keywords = dict(factory.keywords or {})
            keywords["global_config"] = global_config
            return partial(factory.func, *factory.args, **keywords)
        return partial(factory, global_config=global_config)

    def _input_dir_for_workspace(self, workspace_id: str) -> Path:
        current_workspace = self._current_workspace_id()
        current_input_dir = Path(getattr(self.doc_manager, "input_dir", "inputs"))
        if workspace_id == current_workspace:
            return current_input_dir
        base_input_dir = getattr(self.doc_manager, "base_input_dir", None)
        if base_input_dir is not None:
            return Path(base_input_dir) / workspace_id
        return current_input_dir / workspace_id

    async def _workspace_tenant(self, workspace_id: str | None) -> str | None:
        if workspace_id is None:
            return None
        workspace = await self.repository.get_workspace(workspace_id)
        return workspace.tenant_id if workspace else None

    async def _scope_for_workspace(self, workspace_id: str) -> tuple[str | None, str]:
        await self._require_data_plane(workspace_id)
        return await self._workspace_tenant(workspace_id), workspace_id

    async def _existing_workspace_scope(self, workspace_id: str) -> tuple[str | None, str]:
        workspace = await self.repository.get_workspace(workspace_id)
        if workspace is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Workspace '{workspace_id}' not found.",
            )
        return workspace.tenant_id, workspace.workspace_id

    async def _require_existing_group(
        self, *, tenant_id: str | None, workspace_id: str, group_id: str
    ) -> dict[str, Any]:
        if not hasattr(self.admin_store, "get_knowledge_group"):
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Knowledge group registry is unavailable.",
            )
        group = await self.admin_store.get_knowledge_group(
            group_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if not group:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Knowledge group '{group_id}' was not found in workspace '{workspace_id}'.",
            )
        return group

    async def _require_existing_subgroup(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        group_id: str,
        subgroup_id: str,
    ) -> dict[str, Any]:
        if not hasattr(self.admin_store, "get_knowledge_subgroup"):
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Knowledge subgroup registry is unavailable.",
            )
        subgroup = await self.admin_store.get_knowledge_subgroup(
            subgroup_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=group_id,
        )
        if not subgroup:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=(
                    f"Knowledge subgroup '{subgroup_id}' was not found under group "
                    f"'{group_id}' in workspace '{workspace_id}'."
                ),
            )
        return subgroup

    def _current_workspace_id(self) -> str:
        return str(getattr(self.rag, "workspace", None) or "default")

    def _is_backed_workspace(self, workspace_id: str) -> bool:
        return workspace_id == self._current_workspace_id() or workspace_id in self._workspace_rag_cache()

    async def _model_settings_for_workspace(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
    ) -> list[dict[str, Any]]:
        try:
            settings = await self.admin_store.list_model_settings(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
        except Exception:
            settings = []
        if not settings and await self._data_plane_attached(workspace_id):
            settings = self._runtime_model_defaults(workspace_id=workspace_id)
        return settings

    @staticmethod
    def _default_model(settings: list[dict[str, Any]], usage: str) -> dict[str, Any] | None:
        scoped = [
            setting
            for setting in settings
            if setting.get("usage") == usage and setting.get("enabled", True)
        ]
        return next((setting for setting in scoped if setting.get("is_default")), scoped[0] if scoped else None)

    async def _status_counts_for_workspace(self, workspace_id: str) -> dict[str, int]:
        if not await self._data_plane_attached(workspace_id):
            return {}
        rag = await self._require_data_plane(workspace_id)
        return await self._status_counts_safe(rag=rag)

    async def _estimated_tokens_for_workspace(self, workspace_id: str) -> int:
        if not await self._data_plane_attached(workspace_id):
            return 0
        try:
            rag = await self._require_data_plane(workspace_id)
            (documents_with_ids, _total), _counts = await self._documents_paginated(
                rag=rag,
                page=1,
                page_size=200,
            )
        except Exception:
            return 0
        del _total, _counts
        total_characters = 0
        for _doc_id, doc in documents_with_ids:
            total_characters += int(getattr(doc, "content_length", 0) or 0)
        return estimate_tokens_from_characters(total_characters)

    async def list_areas(self, principal: Principal) -> list[LittleBullArea]:
        self._require(principal, ACTIVITY_AREA_READ)
        workspaces = await self.repository.list_workspaces(
            None if principal.is_master_global else principal.tenant_id
        )
        if not principal.is_master_global:
            workspaces = [
                workspace for workspace in workspaces if workspace.workspace_id in principal.workspace_ids
            ]
        areas: list[LittleBullArea] = []
        for workspace in workspaces:
            counts = await self._status_counts_for_workspace(workspace.workspace_id)
            settings = await self._model_settings_for_workspace(
                tenant_id=workspace.tenant_id,
                workspace_id=workspace.workspace_id,
            )
            chat_model = self._default_model(settings, "chat")
            embedding_model = self._default_model(settings, "embedding")
            label, emoji, accent = AREA_PRESENTATION.get(
                workspace.slug,
                (workspace.name, "📁", "#2563EB"),
            )
            areas.append(
                LittleBullArea(
                    id=workspace.workspace_id,
                    label=label or workspace.name,
                    slug=workspace.slug,
                    description=workspace.description,
                    privacy=workspace.privacy,
                    document_count=sum(counts.values()),
                    ready_count=counts.get("processed", 0),
                    processing_count=counts.get("processing", 0) + counts.get("pending", 0),
                    accent=accent,
                    emoji=emoji,
                    data_plane_attached=await self._data_plane_attached(workspace.workspace_id),
                    chat_model_id=chat_model.get("model_id") if chat_model else None,
                    embedding_model_id=embedding_model.get("model_id") if embedding_model else None,
                    embedding_reindex_required=bool(
                        (embedding_model or {}).get("config", {}).get("reindex_required")
                    ),
                )
            )
        return areas

    async def list_knowledge_groups(self, principal: Principal, *, workspace_id: str) -> list[KnowledgeGroup]:
        self._require(principal, ACTIVITY_AREA_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        groups = await self.admin_store.list_knowledge_groups(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        return [KnowledgeGroup(**group) for group in groups]

    async def upsert_knowledge_group(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullKnowledgeGroupRequest,
    ) -> KnowledgeGroup:
        self._require(principal, ACTIVITY_WORKSPACE_MANAGE, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        slug = slugify_workspace(payload.slug or payload.name)
        row = await self.admin_store.upsert_knowledge_group(
            {
                **payload.model_dump(exclude_none=True),
                "slug": slug,
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_WORKSPACE_MANAGE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="knowledge_group_upserted",
            metadata={"group_id": row["group_id"], "slug": row["slug"]},
        )
        return KnowledgeGroup(**row)

    async def list_knowledge_subgroups(
        self, principal: Principal, *, workspace_id: str, group_id: str | None = None
    ) -> list[KnowledgeSubgroup]:
        self._require(principal, ACTIVITY_AREA_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        if group_id:
            await self._require_existing_group(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
            )
        subgroups = await self.admin_store.list_knowledge_subgroups(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=group_id,
        )
        return [KnowledgeSubgroup(**subgroup) for subgroup in subgroups]

    async def upsert_knowledge_subgroup(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullKnowledgeSubgroupRequest,
    ) -> KnowledgeSubgroup:
        self._require(principal, ACTIVITY_WORKSPACE_MANAGE, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        await self._require_existing_group(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
        )
        slug = slugify_workspace(payload.slug or payload.name)
        row = await self.admin_store.upsert_knowledge_subgroup(
            {
                **payload.model_dump(exclude_none=True),
                "slug": slug,
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_WORKSPACE_MANAGE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="knowledge_subgroup_upserted",
            metadata={
                "group_id": row["group_id"],
                "subgroup_id": row["subgroup_id"],
                "slug": row["slug"],
            },
        )
        return KnowledgeSubgroup(**row)

    async def list_notes(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        group_id: str | None = None,
        subgroup_id: str | None = None,
    ) -> list[NoteRegistry]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        if group_id:
            await self._require_existing_group(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
            )
        if subgroup_id:
            await self._require_existing_subgroup(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
                subgroup_id=subgroup_id,
            )
        rows = await self.admin_store.list_note_registry(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=group_id,
            subgroup_id=subgroup_id,
        )
        return [NoteRegistry(**row) for row in rows]

    async def list_tags(self, principal: Principal, *, workspace_id: str) -> list[TagRegistry]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        rows = await self.admin_store.list_tag_registry(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        return [TagRegistry(**row) for row in rows]

    async def _require_workspace_ref(
        self,
        *,
        ref_kind: str,
        ref_id: str,
        tenant_id: str | None,
        workspace_id: str,
    ) -> dict[str, Any] | None:
        normalized_kind = canonical_ref_kind(ref_kind)
        if normalized_kind == "note":
            note = await self.admin_store.get_note_registry(
                ref_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not note:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Note reference not found.")
            return note
        if normalized_kind == "document":
            document = await self.admin_store.get_document_registry(
                ref_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not document:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document reference not found.")
            return document
        if normalized_kind in {"note_label", "answer", "conversation", "agent", "canvas", "trail", "content_map"}:
            return None
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Unsupported reference kind.")

    @staticmethod
    def _refs_share_group_scope(left: dict[str, Any] | None, right: dict[str, Any] | None) -> bool:
        if not left or not right:
            return True
        if left.get("group_id") and right.get("group_id") and left.get("group_id") != right.get("group_id"):
            return False
        if (
            left.get("subgroup_id")
            and right.get("subgroup_id")
            and left.get("subgroup_id") != right.get("subgroup_id")
        ):
            return False
        return True

    async def list_backlinks(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        source_kind: str | None = None,
        source_id: str | None = None,
        target_kind: str | None = None,
        target_id: str | None = None,
    ) -> list[Backlink]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        rows = await self.admin_store.list_backlinks(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            source_kind=source_kind,
            source_id=source_id,
            target_kind=target_kind,
            target_id=target_id,
        )
        return [Backlink(**row) for row in rows]

    async def upsert_backlink(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullBacklinkRequest,
    ) -> Backlink:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        if payload.graph_edge_origin_id:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="graph_edge_origin_id backlinks require scoped graph edge origin validation.",
            )
        source_ref = await self._require_workspace_ref(
            ref_kind=payload.source_kind,
            ref_id=payload.source_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        target_ref = await self._require_workspace_ref(
            ref_kind=payload.target_kind,
            ref_id=payload.target_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if not self._refs_share_group_scope(source_ref, target_ref):
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Backlink source and target must share group/subgroup scope.",
            )
        backlink_payload = payload.model_dump(exclude_none=True)
        backlink_payload["source_kind"] = canonical_ref_kind(payload.source_kind)
        backlink_payload["target_kind"] = canonical_ref_kind(payload.target_kind)
        row = await self.admin_store.upsert_backlink(
            backlink_payload,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="backlink_upserted",
            metadata={
                "backlink_id": row["backlink_id"],
                "source_kind": row["source_kind"],
                "source_id": row["source_id"],
                "target_kind": row["target_kind"],
                "target_id": row["target_id"],
                "origin_type": row["origin_type"],
            },
        )
        return Backlink(**row)

    async def list_source_provenance(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        source_kind: str | None = None,
        source_id: str | None = None,
        document_id: str | None = None,
        note_id: str | None = None,
    ) -> list[SourceProvenance]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        rows = await self.admin_store.list_source_provenance(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            source_kind=source_kind,
            source_id=source_id,
            document_id=document_id,
            note_id=note_id,
        )
        return [SourceProvenance(**row) for row in rows]

    async def record_source_provenance(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullSourceProvenanceRequest,
    ) -> SourceProvenance:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        document_ref = None
        note_ref = None
        if payload.document_id:
            document_ref = await self._require_workspace_ref(
                ref_kind="document",
                ref_id=payload.document_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
        if payload.note_id:
            note_ref = await self._require_workspace_ref(
                ref_kind="note",
                ref_id=payload.note_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
        if not self._refs_share_group_scope(document_ref, note_ref):
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Source provenance references must share group/subgroup scope.",
            )
        if payload.agent_id or payload.usage_ledger_id:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="agent_id and usage_ledger_id provenance require scoped validation before use.",
            )
        row = await self.admin_store.insert_source_provenance(
            payload.model_dump(exclude_none=True),
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="source_provenance_recorded",
            metadata={
                "source_provenance_id": row["source_provenance_id"],
                "source_kind": row["source_kind"],
                "source_id": row["source_id"],
                "document_id": row["document_id"],
                "note_id": row["note_id"],
            },
        )
        return SourceProvenance(**row)

    async def get_provenance_panel(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        target_kind: str,
        target_id: str,
    ) -> LittleBullProvenancePanel:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        canonical_target_kind = canonical_ref_kind(target_kind)
        if canonical_target_kind not in {"note", "document"}:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Provenance panels are currently scoped to note or document targets.",
            )
        await self._require_workspace_ref(
            ref_kind=canonical_target_kind,
            ref_id=target_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        backlinks = await self.admin_store.list_backlinks(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            target_kind=canonical_target_kind,
            target_id=target_id,
        )
        cited_by = [
            row
            for row in backlinks
            if row.get("origin_type") in {"citation", "wikilink", "manual", "source_provenance"}
        ]
        provenance = await self.admin_store.list_source_provenance(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            document_id=target_id if canonical_target_kind == "document" else None,
            note_id=target_id if canonical_target_kind == "note" else None,
        )
        used_in_responses = [
            row
            for row in provenance
            if row.get("source_kind") in {"answer", "query_response", "conversation_message"}
        ]
        return LittleBullProvenancePanel(
            target_kind=canonical_target_kind,
            target_id=target_id,
            mentioned_in=[Backlink(**row) for row in backlinks],
            cited_by=[Backlink(**row) for row in cited_by],
            used_in_responses=[SourceProvenance(**row) for row in used_in_responses],
        )

    async def list_canvas_boards(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        group_id: str | None = None,
        subgroup_id: str | None = None,
    ) -> list[CanvasBoard]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        if group_id:
            await self._require_existing_group(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
            )
        if subgroup_id:
            await self._require_existing_subgroup(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
                subgroup_id=subgroup_id,
            )
        rows = await self.admin_store.list_canvas_boards(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=group_id,
            subgroup_id=subgroup_id,
        )
        return [CanvasBoard(**row) for row in rows]

    async def upsert_canvas_board(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullCanvasBoardRequest,
    ) -> CanvasBoard:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        await self._require_existing_group(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
        )
        await self._require_existing_subgroup(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
            subgroup_id=payload.subgroup_id,
        )
        slug = slugify_workspace(payload.slug or payload.title)
        if payload.canvas_board_id:
            existing_by_id = await self.admin_store.get_canvas_board(
                payload.canvas_board_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not existing_by_id:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Canvas board not found.")
        existing_boards = await self.admin_store.list_canvas_boards(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        found_canvas_board_id = False
        for existing in existing_boards:
            if existing["canvas_board_id"] == payload.canvas_board_id:
                found_canvas_board_id = True
            if payload.canvas_board_id and existing["slug"] == slug and existing["canvas_board_id"] != payload.canvas_board_id:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail="Canvas board slug is already used by another board.",
                )
            if existing["slug"] == slug or existing["canvas_board_id"] == payload.canvas_board_id:
                if existing["group_id"] != payload.group_id or existing["subgroup_id"] != payload.subgroup_id:
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                        detail="Existing canvas boards cannot be moved across group/subgroup scope.",
                    )
        if payload.canvas_board_id and not found_canvas_board_id:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Canvas board not found.")
        try:
            row = await self.admin_store.upsert_canvas_board(
                {
                    **payload.model_dump(exclude_none=True),
                    "slug": slug,
                },
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                user_id=principal.user_id,
            )
        except ValueError as exc:
            if "canvas_board_scope_mismatch" not in str(exc):
                raise
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Existing canvas boards cannot be moved across group/subgroup scope.",
            ) from exc
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="canvas_board_upserted",
            metadata={"canvas_board_id": row["canvas_board_id"], "slug": row["slug"]},
        )
        return CanvasBoard(**row)

    async def _require_canvas_board(
        self,
        *,
        canvas_board_id: str,
        tenant_id: str | None,
        workspace_id: str,
    ) -> dict[str, Any]:
        board = await self.admin_store.get_canvas_board(
            canvas_board_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if not board:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Canvas board not found.")
        return board

    async def get_canvas_board(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        canvas_board_id: str,
    ) -> LittleBullCanvasBoardDetail:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        board = await self._require_canvas_board(
            canvas_board_id=canvas_board_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        nodes = await self.admin_store.list_canvas_nodes(
            canvas_board_id=canvas_board_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        edges = await self.admin_store.list_canvas_edges(
            canvas_board_id=canvas_board_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        return LittleBullCanvasBoardDetail(
            board=CanvasBoard(**board),
            nodes=[CanvasNode(**node) for node in nodes],
            edges=[CanvasEdge(**edge) for edge in edges],
        )

    async def upsert_canvas_node(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        canvas_board_id: str,
        payload: LittleBullCanvasNodeRequest,
    ) -> CanvasNode:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        board = await self._require_canvas_board(
            canvas_board_id=canvas_board_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        ref_kind = canonical_ref_kind(payload.ref_kind) if payload.ref_kind else ""
        if ref_kind in {"note", "document"}:
            ref = await self._require_workspace_ref(
                ref_kind=ref_kind,
                ref_id=payload.ref_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not self._refs_share_group_scope(board, ref):
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Canvas node reference must share board group/subgroup scope.",
                )
        elif payload.ref_id:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Canvas nodes with ref_id currently support only note or document refs.",
            )
        if payload.canvas_node_id:
            existing_node = await self.admin_store.get_canvas_node(
                payload.canvas_node_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not existing_node or existing_node["canvas_board_id"] != canvas_board_id:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Canvas node id is not available on this board.",
                )
        row = await self.admin_store.upsert_canvas_node(
            {
                **payload.model_dump(exclude_none=True),
                "canvas_board_id": canvas_board_id,
                "ref_kind": ref_kind,
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="canvas_node_upserted",
            metadata={"canvas_board_id": canvas_board_id, "canvas_node_id": row["canvas_node_id"]},
        )
        return CanvasNode(**row)

    async def upsert_canvas_edge(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        canvas_board_id: str,
        payload: LittleBullCanvasEdgeRequest,
    ) -> CanvasEdge:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        await self._require_canvas_board(
            canvas_board_id=canvas_board_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        source = await self.admin_store.get_canvas_node(
            payload.source_node_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        target = await self.admin_store.get_canvas_node(
            payload.target_node_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if not source or not target or source["canvas_board_id"] != canvas_board_id or target["canvas_board_id"] != canvas_board_id:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Canvas edge endpoints must be nodes on the same board.",
            )
        if payload.canvas_edge_id:
            existing_edges = await self.admin_store.list_canvas_edges(
                canvas_board_id=canvas_board_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not any(edge["canvas_edge_id"] == payload.canvas_edge_id for edge in existing_edges):
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Canvas edge id is not available on this board.",
                )
        row = await self.admin_store.upsert_canvas_edge(
            {
                **payload.model_dump(exclude_none=True),
                "canvas_board_id": canvas_board_id,
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="canvas_edge_upserted",
            metadata={"canvas_board_id": canvas_board_id, "canvas_edge_id": row["canvas_edge_id"]},
        )
        return CanvasEdge(**row)

    async def analyze_canvas_board(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        canvas_board_id: str,
    ) -> LittleBullCanvasAnalysis:
        detail = await self.get_canvas_board(principal, workspace_id=workspace_id, canvas_board_id=canvas_board_id)
        node_kind_counts: dict[str, int] = {}
        for node in detail.nodes:
            node_kind_counts[node.node_kind] = node_kind_counts.get(node.node_kind, 0) + 1

        adjacency: dict[str, set[str]] = {node.canvas_node_id or "": set() for node in detail.nodes}
        for edge in detail.edges:
            adjacency.setdefault(edge.source_node_id, set()).add(edge.target_node_id)
            adjacency.setdefault(edge.target_node_id, set()).add(edge.source_node_id)
        clusters = []
        seen: set[str] = set()
        for node_id in adjacency:
            if not node_id or node_id in seen:
                continue
            stack = [node_id]
            cluster_nodes: list[str] = []
            seen.add(node_id)
            while stack:
                current = stack.pop()
                cluster_nodes.append(current)
                for neighbor in adjacency.get(current, set()):
                    if neighbor not in seen:
                        seen.add(neighbor)
                        stack.append(neighbor)
            clusters.append({"cluster_id": f"cluster-{len(clusters) + 1}", "node_ids": sorted(cluster_nodes)})
        warnings = []
        if not detail.nodes:
            warnings.append("canvas_empty")
        if detail.nodes and not detail.edges:
            warnings.append("canvas_without_edges")
        return LittleBullCanvasAnalysis(
            canvas_board_id=canvas_board_id,
            node_count=len(detail.nodes),
            edge_count=len(detail.edges),
            node_kind_counts=node_kind_counts,
            clusters=clusters,
            warnings=warnings,
        )

    async def export_canvas_board_dossier(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        canvas_board_id: str,
    ) -> KnowledgeDossier:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        detail = await self.get_canvas_board(principal, workspace_id=workspace_id, canvas_board_id=canvas_board_id)
        analysis = await self.analyze_canvas_board(principal, workspace_id=workspace_id, canvas_board_id=canvas_board_id)
        row = await self.admin_store.upsert_knowledge_dossier(
            {
                "group_id": detail.board.group_id,
                "subgroup_id": detail.board.subgroup_id,
                "title": f"{detail.board.title} dossier",
                "slug": slugify_workspace(f"{detail.board.slug}-dossier"),
                "dossier_kind": "canvas",
                "status": "draft",
                "content_refs": [
                    {"kind": "canvas_board", "id": canvas_board_id},
                    *[
                        {"kind": node.node_kind, "id": node.ref_id or node.canvas_node_id}
                        for node in detail.nodes
                    ],
                ],
                "export_policy": {
                    "requires_lgpd_review": True,
                    "source": "canvas",
                    "analysis": analysis.model_dump(),
                },
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="canvas_dossier_exported",
            metadata={"canvas_board_id": canvas_board_id, "knowledge_dossier_id": row["knowledge_dossier_id"]},
        )
        return KnowledgeDossier(**row)

    async def list_content_maps(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        group_id: str | None = None,
        subgroup_id: str | None = None,
    ) -> list[ContentMap]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        if group_id:
            await self._require_existing_group(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
            )
        if subgroup_id:
            await self._require_existing_subgroup(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
                subgroup_id=subgroup_id,
            )
        rows = await self.admin_store.list_content_maps(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=group_id,
            subgroup_id=subgroup_id,
        )
        return [ContentMap(**row) for row in rows]

    async def upsert_content_map(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullContentMapRequest,
    ) -> ContentMap:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        await self._require_existing_group(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
        )
        await self._require_existing_subgroup(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
            subgroup_id=payload.subgroup_id,
        )
        if payload.root_note_id:
            root_note = await self._require_workspace_ref(
                ref_kind="note",
                ref_id=payload.root_note_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            content_map_scope = {
                "group_id": payload.group_id,
                "subgroup_id": payload.subgroup_id,
            }
            if not self._refs_share_group_scope(content_map_scope, root_note):
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Content map root note must share group/subgroup scope.",
                )
        slug = slugify_workspace(payload.slug or payload.title)
        existing_maps = await self.admin_store.list_content_maps(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        found_content_map_id = False
        for existing in existing_maps:
            if existing["content_map_id"] == payload.content_map_id:
                found_content_map_id = True
            if payload.content_map_id and existing["slug"] == slug and existing["content_map_id"] != payload.content_map_id:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail="Content map slug is already used by another content map.",
                )
            if existing["slug"] == slug or existing["content_map_id"] == payload.content_map_id:
                if existing["group_id"] != payload.group_id or existing["subgroup_id"] != payload.subgroup_id:
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                        detail="Existing content maps cannot be moved across group/subgroup scope.",
                    )
        if payload.content_map_id and not found_content_map_id:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Content map not found.")
        try:
            row = await self.admin_store.upsert_content_map(
                {
                    **payload.model_dump(exclude_none=True),
                    "slug": slug,
                },
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                user_id=principal.user_id,
            )
        except ValueError as exc:
            if "content_map_scope_mismatch" not in str(exc):
                raise
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Existing content maps cannot be moved across group/subgroup scope.",
            ) from exc
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="content_map_upserted",
            metadata={"content_map_id": row["content_map_id"], "slug": row["slug"]},
        )
        return ContentMap(**row)

    async def list_knowledge_trails(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        group_id: str | None = None,
        subgroup_id: str | None = None,
    ) -> list[KnowledgeTrail]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        if group_id:
            await self._require_existing_group(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
            )
        if subgroup_id:
            await self._require_existing_subgroup(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
                subgroup_id=subgroup_id,
            )
        rows = await self.admin_store.list_knowledge_trails(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=group_id,
            subgroup_id=subgroup_id,
        )
        return [KnowledgeTrail(**row) for row in rows]

    async def upsert_knowledge_trail(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullKnowledgeTrailRequest,
    ) -> KnowledgeTrail:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        await self._require_existing_group(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
        )
        await self._require_existing_subgroup(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
            subgroup_id=payload.subgroup_id,
        )
        slug = slugify_workspace(payload.slug or payload.title)
        existing_trails = await self.admin_store.list_knowledge_trails(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        found_knowledge_trail_id = False
        for existing in existing_trails:
            if existing["knowledge_trail_id"] == payload.knowledge_trail_id:
                found_knowledge_trail_id = True
            if (
                payload.knowledge_trail_id
                and existing["slug"] == slug
                and existing["knowledge_trail_id"] != payload.knowledge_trail_id
            ):
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail="Knowledge trail slug is already used by another trail.",
                )
            if existing["slug"] == slug or existing["knowledge_trail_id"] == payload.knowledge_trail_id:
                if existing["group_id"] != payload.group_id or existing["subgroup_id"] != payload.subgroup_id:
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                        detail="Existing knowledge trails cannot be moved across group/subgroup scope.",
                    )
        if payload.knowledge_trail_id and not found_knowledge_trail_id:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Knowledge trail not found.")
        try:
            row = await self.admin_store.upsert_knowledge_trail(
                {
                    **payload.model_dump(exclude_none=True),
                    "slug": slug,
                },
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                user_id=principal.user_id,
            )
        except ValueError as exc:
            if "knowledge_trail_scope_mismatch" not in str(exc):
                raise
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Existing knowledge trails cannot be moved across group/subgroup scope.",
            ) from exc
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="knowledge_trail_upserted",
            metadata={"knowledge_trail_id": row["knowledge_trail_id"], "slug": row["slug"]},
        )
        return KnowledgeTrail(**row)

    async def _require_knowledge_trail(
        self,
        *,
        knowledge_trail_id: str,
        tenant_id: str | None,
        workspace_id: str,
    ) -> dict[str, Any]:
        trail = await self.admin_store.get_knowledge_trail(
            knowledge_trail_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if not trail:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Knowledge trail not found.")
        return trail

    async def get_knowledge_trail(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        knowledge_trail_id: str,
    ) -> LittleBullKnowledgeTrailDetail:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        trail = await self._require_knowledge_trail(
            knowledge_trail_id=knowledge_trail_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        steps = await self.admin_store.list_knowledge_trail_steps(
            knowledge_trail_id=knowledge_trail_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        return LittleBullKnowledgeTrailDetail(
            trail=KnowledgeTrail(**trail),
            steps=[KnowledgeTrailStep(**step) for step in steps],
        )

    async def upsert_knowledge_trail_step(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        knowledge_trail_id: str,
        payload: LittleBullKnowledgeTrailStepRequest,
    ) -> KnowledgeTrailStep:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        trail = await self._require_knowledge_trail(
            knowledge_trail_id=knowledge_trail_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        existing_steps = await self.admin_store.list_knowledge_trail_steps(
            knowledge_trail_id=knowledge_trail_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if payload.knowledge_trail_step_id and not any(
            step["knowledge_trail_step_id"] == payload.knowledge_trail_step_id for step in existing_steps
        ):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Knowledge trail step id is not available on this trail.",
            )
        trail_scope = {"group_id": trail.get("group_id"), "subgroup_id": trail.get("subgroup_id")}
        for ref_kind, ref_id in (
            ("note", payload.note_id),
            ("document", payload.document_id),
            ("canvas", payload.canvas_board_id),
        ):
            if not ref_id:
                continue
            if ref_kind == "canvas":
                ref = await self._require_canvas_board(
                    canvas_board_id=ref_id,
                    tenant_id=tenant_id,
                    workspace_id=workspace_id,
                )
            else:
                ref = await self._require_workspace_ref(
                    ref_kind=ref_kind,
                    ref_id=ref_id,
                    tenant_id=tenant_id,
                    workspace_id=workspace_id,
                )
            if not self._refs_share_group_scope(trail_scope, ref):
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Knowledge trail step references must share trail group/subgroup scope.",
                )
        row = await self.admin_store.upsert_knowledge_trail_step(
            {
                **payload.model_dump(exclude_none=True),
                "knowledge_trail_id": knowledge_trail_id,
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="knowledge_trail_step_upserted",
            metadata={
                "knowledge_trail_id": knowledge_trail_id,
                "knowledge_trail_step_id": row["knowledge_trail_step_id"],
                "step_order": row["step_order"],
            },
        )
        return KnowledgeTrailStep(**row)

    @staticmethod
    def _parse_daily_note_date(value: str | None) -> str:
        if not value:
            return datetime.now(timezone.utc).date().isoformat()
        try:
            return date.fromisoformat(value).isoformat()
        except ValueError as exc:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="note_date must be an ISO date in YYYY-MM-DD format.",
            ) from exc

    async def _require_optional_group_scope(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        group_id: str | None,
        subgroup_id: str | None,
    ) -> dict[str, Any]:
        if subgroup_id and not group_id:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="subgroup_id requires group_id.",
            )
        if group_id:
            await self._require_existing_group(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
            )
        if subgroup_id:
            await self._require_existing_subgroup(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                group_id=group_id,
                subgroup_id=subgroup_id,
            )
        return {"group_id": group_id, "subgroup_id": subgroup_id}

    async def _require_content_map(
        self,
        *,
        content_map_id: str,
        tenant_id: str | None,
        workspace_id: str,
    ) -> dict[str, Any]:
        content_map = next(
            (
                item
                for item in await self.admin_store.list_content_maps(
                    tenant_id=tenant_id,
                    workspace_id=workspace_id,
                )
                if item["content_map_id"] == content_map_id
            ),
            None,
        )
        if not content_map:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Content map not found.")
        return content_map

    async def _validate_inbox_source(
        self,
        *,
        principal: Principal,
        tenant_id: str | None,
        workspace_id: str,
        group_scope: dict[str, Any],
        source_kind: str,
        source_id: str,
    ) -> str:
        source_kind = canonical_ref_kind(source_kind)
        if bool(source_kind) != bool(source_id):
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Inbox source_kind and source_id must be provided together.",
            )
        if not source_kind:
            return ""
        if source_kind in {"note", "document"}:
            ref = await self._require_workspace_ref(
                ref_kind=source_kind,
                ref_id=source_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
        elif source_kind == "canvas":
            ref = await self._require_canvas_board(
                canvas_board_id=source_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
        elif source_kind == "trail":
            ref = await self._require_knowledge_trail(
                knowledge_trail_id=source_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
        elif source_kind == "content_map":
            ref = await self._require_content_map(
                content_map_id=source_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
        elif source_kind == "conversation":
            ref = await self.admin_store.get_conversation(source_id)
            if not ref or ref["tenant_id"] != tenant_id or ref["workspace_id"] != workspace_id:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Conversation source not found.")
            if not principal.is_master_global and ref["user_id"] != principal.user_id:
                raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Conversation belongs to another user.")
            return source_kind
        elif source_kind == "suggestion":
            ref = await self.admin_store.get_correlation_suggestion(source_id)
            if not ref or ref["tenant_id"] != tenant_id or ref["workspace_id"] != workspace_id:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Suggestion source not found.")
            if not principal.is_master_global and ref["user_id"] != principal.user_id:
                raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Suggestion belongs to another user.")
            return source_kind
        else:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Unsupported inbox source.")
        if (ref.get("group_id") or ref.get("subgroup_id")) and not (
            group_scope.get("group_id") and group_scope.get("subgroup_id")
        ):
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Scoped inbox sources require inbox group/subgroup scope.",
            )
        if not self._refs_share_group_scope(group_scope, ref):
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Inbox source must share inbox group/subgroup scope.",
            )
        return source_kind

    async def list_inbox_items(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        status_filter: str | None = None,
        group_id: str | None = None,
        subgroup_id: str | None = None,
        limit: int = 100,
    ) -> list[KnowledgeInboxItem]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        await self._require_optional_group_scope(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=group_id,
            subgroup_id=subgroup_id,
        )
        rows = await self.admin_store.list_knowledge_inbox_items(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            status=status_filter,
            group_id=group_id,
            subgroup_id=subgroup_id,
            limit=limit,
        )
        return [KnowledgeInboxItem(**row) for row in rows]

    async def upsert_inbox_item(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullInboxItemRequest,
    ) -> KnowledgeInboxItem:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        group_scope = await self._require_optional_group_scope(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
            subgroup_id=payload.subgroup_id,
        )
        source_kind = await self._validate_inbox_source(
            principal=principal,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_scope=group_scope,
            source_kind=payload.source_kind,
            source_id=payload.source_id,
        )
        if payload.inbox_item_id:
            existing = await self.admin_store.get_knowledge_inbox_item(
                payload.inbox_item_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not existing:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Inbox item not found.")
            if existing["group_id"] != payload.group_id or existing["subgroup_id"] != payload.subgroup_id:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Existing inbox items cannot be moved across group/subgroup scope.",
                )
        row = await self.admin_store.upsert_knowledge_inbox_item(
            {
                **payload.model_dump(exclude_none=True),
                "source_kind": source_kind,
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="inbox_item_upserted",
            metadata={"inbox_item_id": row["inbox_item_id"], "item_kind": row["item_kind"]},
        )
        return KnowledgeInboxItem(**row)

    async def update_inbox_item_status(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        inbox_item_id: str,
        payload: LittleBullInboxItemStatusRequest,
    ) -> KnowledgeInboxItem:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        row = await self.admin_store.update_knowledge_inbox_item_status(
            inbox_item_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            status=payload.status.strip(),
            metadata=payload.metadata,
            user_id=principal.user_id,
        )
        if not row:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Inbox item not found.")
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="inbox_item_status_updated",
            metadata={"inbox_item_id": inbox_item_id, "status": row["status"]},
        )
        return KnowledgeInboxItem(**row)

    @staticmethod
    def _daily_note_markdown(note_date: str, payload: LittleBullDailyNoteRequest) -> str:
        decisions = payload.decisions or []
        pending_items = payload.pending_items or []
        decision_lines = "\n".join(f"- {item.get('title') or item}" for item in decisions) or "- Nenhuma decisão registrada."
        pending_lines = "\n".join(f"- {item.get('title') or item}" for item in pending_items) or "- Nenhuma pendência aberta."
        cost_lines = "\n".join(f"- {key}: {value}" for key, value in sorted(payload.cost_snapshot.items())) or "- Sem custos registrados."
        summary = payload.summary.strip() or "Sem resumo registrado."
        return (
            f"# Daily Note {note_date}\n\n"
            f"## Resumo\n{summary}\n\n"
            f"## Decisões\n{decision_lines}\n\n"
            f"## Pendências\n{pending_lines}\n\n"
            f"## Custos\n{cost_lines}\n"
        )

    async def list_daily_notes(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        limit: int = 30,
    ) -> list[DailyNote]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        rows = await self.admin_store.list_daily_notes(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            limit=limit,
        )
        return [DailyNote(**row) for row in rows]

    async def ensure_daily_note(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullDailyNoteRequest,
    ) -> DailyNote:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        await self._require_existing_group(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
        )
        await self._require_existing_subgroup(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
            subgroup_id=payload.subgroup_id,
        )
        note_date = self._parse_daily_note_date(payload.note_date)
        existing = await self.admin_store.get_daily_note(
            note_date,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        note_id = existing["note_id"] if existing else None
        daily_slug = slugify_workspace(f"daily-{note_date}")
        slug_note = await self.admin_store.find_note_by_slug_or_title(
            slug=daily_slug,
            title=f"Daily Note {note_date}",
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if slug_note and slug_note["note_id"] != note_id:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="Daily note slug is already used by another note.",
            )
        if existing:
            registry = await self.admin_store.get_note_registry(
                note_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not registry or registry["group_id"] != payload.group_id or registry["subgroup_id"] != payload.subgroup_id:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Existing daily notes cannot be moved across group/subgroup scope.",
                )
        if not payload.pending_items:
            open_items = await self.admin_store.list_knowledge_inbox_items(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                status="open",
                group_id=payload.group_id,
                subgroup_id=payload.subgroup_id,
                limit=25,
            )
            payload.pending_items.extend(
                {"inbox_item_id": item["inbox_item_id"], "title": item["title"], "priority": item["priority"]}
                for item in open_items
            )
        markdown_note = await self.upsert_markdown_note(
            principal,
            workspace_id=workspace_id,
            payload=LittleBullMarkdownNoteRequest(
                note_id=note_id,
                title=f"Daily Note {note_date}",
                slug=daily_slug,
                group_id=payload.group_id,
                subgroup_id=payload.subgroup_id,
                markdown=self._daily_note_markdown(note_date, payload),
                metadata={"daily_note": True, "note_date": note_date},
            ),
        )
        row = await self.admin_store.upsert_daily_note(
            {
                "daily_note_id": payload.daily_note_id or (existing or {}).get("daily_note_id"),
                "note_id": markdown_note.registry.note_id,
                "note_date": note_date,
                "summary": payload.summary,
                "decisions": payload.decisions,
                "pending_items": payload.pending_items,
                "cost_snapshot": payload.cost_snapshot,
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="daily_note_upserted",
            metadata={"daily_note_id": row["daily_note_id"], "note_id": row["note_id"], "note_date": row["note_date"]},
        )
        return DailyNote(**row)

    async def get_markdown_note(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        note_id: str,
    ) -> LittleBullMarkdownNoteResponse:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        registry = await self.admin_store.get_note_registry(
            note_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if not registry:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Note not found.")
        note = await self.admin_store.get_latest_markdown_note(
            note_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if not note:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Markdown note not found.")
        wiki_links = await self.admin_store.list_wiki_links(
            source_note_id=note_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        note_tags = set(registry.get("metadata", {}).get("tags") or [])
        tags = [
            tag
            for tag in await self.admin_store.list_tag_registry(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if tag.get("tag") in note_tags
        ]
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_READ,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="markdown_note_read",
            metadata={"note_id": note_id},
        )
        return LittleBullMarkdownNoteResponse(
            registry=NoteRegistry(**registry),
            note=MarkdownNote(**note),
            wiki_links=[WikiLink(**link) for link in wiki_links],
            tags=[TagRegistry(**tag) for tag in tags],
        )

    async def upsert_markdown_note(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullMarkdownNoteRequest,
    ) -> LittleBullMarkdownNoteResponse:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        if not payload.group_id or not payload.subgroup_id:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="group_id and subgroup_id are required for markdown notes.",
            )
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        await self._require_existing_group(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
        )
        await self._require_existing_subgroup(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=payload.group_id,
            subgroup_id=payload.subgroup_id,
        )
        if payload.source_document_id:
            document = await self.admin_store.get_document_registry(
                payload.source_document_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not document:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Source document not found.")
            if document.get("group_id") != payload.group_id or document.get("subgroup_id") != payload.subgroup_id:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Source document must belong to the same group and subgroup as the note.",
                )

        wiki_link_specs = extract_wiki_links(payload.markdown)
        markdown_tags = extract_markdown_tags(payload.markdown)
        slug = slugify_workspace(payload.slug or payload.title)
        registry = await self.admin_store.upsert_note_registry(
            {
                "note_id": payload.note_id,
                "group_id": payload.group_id,
                "subgroup_id": payload.subgroup_id,
                "title": payload.title.strip(),
                "slug": slug,
                "note_type": "markdown",
                "privacy": payload.privacy.strip() or "team",
                "status": "active",
                "metadata": {
                    **payload.metadata,
                    "tags": markdown_tags,
                    "wikilinks": [link["target_label"] for link in wiki_link_specs],
                },
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        note = await self.admin_store.insert_markdown_note_version(
            {
                "note_id": registry["note_id"],
                "markdown": payload.markdown,
                "rendered_summary": markdown_summary(payload.markdown),
                "content_hash": f"sha256:{hashlib.sha256(payload.markdown.encode('utf-8')).hexdigest()}",
                "status": "current",
                "source_document_id": payload.source_document_id,
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )

        wiki_payloads: list[dict[str, Any]] = []
        for link in wiki_link_specs:
            target_label = link["target_label"]
            target = await self.admin_store.find_note_by_slug_or_title(
                slug=slugify_workspace(target_label),
                title=target_label,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if target and (
                target.get("group_id") != payload.group_id
                or target.get("subgroup_id") != payload.subgroup_id
            ):
                target = None
            wiki_payloads.append(
                {
                    "source_note_id": registry["note_id"],
                    "target_note_id": target.get("note_id") if target else None,
                    "target_label": target_label,
                    "link_text": link["link_text"],
                    "link_status": "resolved" if target else "unresolved",
                    "metadata": {},
                }
            )
        wiki_links = await self.admin_store.replace_wiki_links(
            source_note_id=registry["note_id"],
            links=wiki_payloads,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.admin_store.replace_backlinks_for_source(
            source_kind="note",
            source_id=registry["note_id"],
            origin_type="wikilink",
            backlinks=[
                {
                    "source_kind": "note",
                    "source_id": registry["note_id"],
                    "target_kind": "note" if link.get("target_note_id") else "note_label",
                    "target_id": link.get("target_note_id") or slugify_workspace(link["target_label"]),
                    "link_text": link.get("link_text", ""),
                    "origin_type": "wikilink",
                    "confidence": 1.0 if link.get("target_note_id") else None,
                    "metadata": {
                        "target_label": link["target_label"],
                        "link_status": link["link_status"],
                    },
                }
                for link in wiki_links
            ],
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )

        tag_rows = []
        for tag in markdown_tags:
            tag_rows.append(
                await self.admin_store.upsert_tag_registry(
                    {
                        "tag": tag,
                        "label": tag.removeprefix("#").replace("-", " ").replace("_", " ").title(),
                        "description": "",
                        "color": "#64748B",
                        "metadata": {},
                    },
                    tenant_id=tenant_id,
                    workspace_id=workspace_id,
                    user_id=principal.user_id,
                )
            )

        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="markdown_note_upserted",
            metadata={
                "note_id": registry["note_id"],
                "group_id": payload.group_id,
                "subgroup_id": payload.subgroup_id,
                "version_number": note["version_number"],
                "wikilink_count": len(wiki_links),
                "tag_count": len(tag_rows),
            },
        )
        return LittleBullMarkdownNoteResponse(
            registry=NoteRegistry(**registry),
            note=MarkdownNote(**note),
            wiki_links=[WikiLink(**link) for link in wiki_links],
            tags=[TagRegistry(**tag) for tag in tag_rows],
        )

    async def list_documents(
        self, principal: Principal, *, workspace_id: str, page: int = 1, page_size: int = 50
    ) -> LittleBullDocumentsResponse:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        tenant_id = await self._workspace_tenant(workspace_id)
        rag = await self._require_data_plane(workspace_id)
        if hasattr(self.admin_store, "list_document_registry"):
            registry_rows = await self.admin_store.list_document_registry(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            classified_rows = [
                row
                for row in registry_rows
                if row.get("group_id") and row.get("subgroup_id")
            ]
            total_count = len(classified_rows)
            start = max(page - 1, 0) * page_size
            page_rows = classified_rows[start : start + page_size]
            status_counts: dict[str, int] = {}
            for row in classified_rows:
                row_status = str(row.get("status") or "registered")
                status_counts[row_status] = status_counts.get(row_status, 0) + 1

            try:
                (status_documents, _ignored_total), _ignored_counts = await self._documents_paginated(
                    rag=rag,
                    page=1,
                    page_size=min(max(page_size, len(page_rows), 200), 200),
                )
            except Exception as exc:
                raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(exc)) from exc

            status_by_source: dict[str, tuple[str, Any]] = {}
            status_by_title: dict[str, tuple[str, Any]] = {}
            for doc_id, doc in status_documents:
                file_path = str(getattr(doc, "file_path", "") or "")
                title = Path(file_path).name or doc_id
                if file_path:
                    status_by_source[file_path] = (doc_id, doc)
                status_by_title[title] = (doc_id, doc)

            documents: list[LittleBullDocument] = []
            for registry in page_rows:
                source_uri = str(registry.get("source_uri") or "")
                source_name = Path(source_uri).name if source_uri else ""
                registry_title = str(registry.get("title") or "")
                title = registry_title or source_name or registry["document_id"]
                status_entry = (
                    status_by_source.get(source_uri)
                    or status_by_title.get(source_name)
                    or status_by_title.get(title)
                )
                doc_id = status_entry[0] if status_entry else registry["document_id"]
                doc = status_entry[1] if status_entry else None
                metadata = dict(getattr(doc, "metadata", {}) or {}) if doc else {}
                metadata = {
                    **metadata,
                    "registry_document_id": registry["document_id"],
                    "group_id": registry["group_id"],
                    "subgroup_id": registry["subgroup_id"],
                    "source_kind": registry["source_kind"],
                }
                documents.append(
                    LittleBullDocument(
                        id=doc_id,
                        file_path=source_name or title,
                        title=title,
                        status=str(getattr(doc, "status", "") or registry.get("status") or "registered"),
                        group_id=registry["group_id"],
                        subgroup_id=registry["subgroup_id"],
                        registry_document_id=registry["document_id"],
                        content_summary=str(getattr(doc, "content_summary", "") or ""),
                        content_length=int(getattr(doc, "content_length", 0) or 0),
                        updated_at=str(getattr(doc, "updated_at", "") or registry.get("updated_at") or "") or None,
                        created_at=str(getattr(doc, "created_at", "") or registry.get("created_at") or "") or None,
                        track_id=getattr(doc, "track_id", None) if doc else registry.get("metadata", {}).get("track_id"),
                        chunks_count=getattr(doc, "chunks_count", None) if doc else registry.get("chunk_count"),
                        metadata=metadata,
                    )
                )
            await self.audit.record(
                principal=principal,
                action=ACTIVITY_DOCUMENT_READ,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                result="success",
                metadata={"total_count": total_count},
            )
            return LittleBullDocumentsResponse(
                documents=documents,
                total_count=total_count,
                status_counts=status_counts,
            )

        try:
            (documents_with_ids, total_count), status_counts = await self._documents_paginated(
                rag=rag,
                page=page,
                page_size=page_size,
            )
        except Exception as exc:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(exc)) from exc

        documents: list[LittleBullDocument] = []
        for doc_id, doc in documents_with_ids:
            file_path = str(getattr(doc, "file_path", "") or "")
            title = Path(file_path).name or doc_id
            metadata = dict(getattr(doc, "metadata", {}) or {})
            documents.append(
                LittleBullDocument(
                    id=doc_id,
                    file_path=file_path,
                    title=title,
                    status=str(getattr(doc, "status", "unknown")),
                    content_summary=str(getattr(doc, "content_summary", "") or ""),
                    content_length=int(getattr(doc, "content_length", 0) or 0),
                    updated_at=str(getattr(doc, "updated_at", "") or "") or None,
                    created_at=str(getattr(doc, "created_at", "") or "") or None,
                    track_id=getattr(doc, "track_id", None),
                    chunks_count=getattr(doc, "chunks_count", None),
                    metadata=metadata,
                )
            )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_READ,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="success",
            metadata={"total_count": total_count},
        )
        return LittleBullDocumentsResponse(
            documents=documents,
            total_count=total_count,
            status_counts=status_counts,
        )

    async def upload_document(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        group_id: str,
        subgroup_id: str,
        file: UploadFile,
        background_tasks: BackgroundTasks,
        confidentiality: str = "normal",
    ) -> LittleBullUploadResponse:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        if not group_id or not subgroup_id:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="group_id and subgroup_id are required for clean uploads.",
            )
        tenant_id = await self._workspace_tenant(workspace_id)
        await self._require_existing_group(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=group_id,
        )
        await self._require_existing_subgroup(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            group_id=group_id,
            subgroup_id=subgroup_id,
        )
        if not hasattr(self.admin_store, "register_document") or not hasattr(self.admin_store, "list_document_registry"):
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Document registry is unavailable.",
            )
        rag = await self._require_data_plane(workspace_id)
        input_dir = self._input_dir_for_workspace(workspace_id)
        input_dir.mkdir(parents=True, exist_ok=True)
        safe_filename = sanitize_upload_filename(file.filename or "document", input_dir)
        if hasattr(self.doc_manager, "is_supported_file") and not self.doc_manager.is_supported_file(safe_filename):
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type")
        registry_rows = await self.admin_store.list_document_registry(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        registered_names = {
            name
            for row in registry_rows
            for name in (
                str(row.get("title") or ""),
                Path(str(row.get("source_uri") or "")).name,
            )
            if name
        }
        safe_filename = unique_input_filename(
            safe_filename,
            input_dir,
            reserved_names=registered_names,
            extra_dirs=[input_dir / "__enqueued__"],
        )
        file_path = input_dir / safe_filename

        content_hash = hashlib.sha256()
        async with aiofiles.open(file_path, "wb") as out_file:
            while chunk := await file.read(1024 * 1024):
                content_hash.update(chunk)
                await out_file.write(chunk)

        if confidentiality in {"sensivel", "privado"}:
            await self.repository.set_policy(
                WORKSPACE_PRIVATE_POLICY,
                True,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )

        track_id = generate_track_id("little_bull_upload")
        registry = await self.admin_store.register_document(
            {
                "group_id": group_id,
                "subgroup_id": subgroup_id,
                "title": safe_filename,
                "source_uri": safe_filename,
                "source_kind": "upload",
                "mime_type": file.content_type or "",
                "content_hash": content_hash.hexdigest(),
                "confidentiality": confidentiality,
                "status": "queued",
                "metadata": {"track_id": track_id},
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        self._queue_pipeline_index_file(background_tasks, file_path, track_id, rag=rag)
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="queued",
            metadata={
                "file_name": safe_filename,
                "track_id": track_id,
                "confidentiality": confidentiality,
                "group_id": group_id,
                "subgroup_id": subgroup_id,
                "registry_document_id": registry["document_id"],
            },
        )
        return LittleBullUploadResponse(
            status="success",
            message=f"File '{safe_filename}' uploaded and queued for indexing.",
            track_id=track_id,
            workspace_id=workspace_id,
            group_id=group_id,
            subgroup_id=subgroup_id,
            registry_document_id=registry["document_id"],
        )

    async def reindex_archived_documents(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        background_tasks: BackgroundTasks,
    ) -> LittleBullReindexArchivedResponse:
        self._require(principal, ACTIVITY_DOCUMENT_UPLOAD, workspace_id)
        rag = await self._require_data_plane(workspace_id)
        input_dir = self._input_dir_for_workspace(workspace_id)
        archived_dir = input_dir / "__enqueued__"
        tenant_id = await self._workspace_tenant(workspace_id)
        if not archived_dir.exists():
            await self.audit.record(
                principal=principal,
                action=ACTIVITY_DOCUMENT_UPLOAD,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                result="no_archived_files",
                metadata={"source": "__enqueued__", "file_count": 0},
            )
            return LittleBullReindexArchivedResponse(
                status="no_files",
                message="No archived files were found for this workspace.",
                workspace_id=workspace_id,
            )

        input_dir.mkdir(parents=True, exist_ok=True)
        copied_paths: list[Path] = []
        skipped: list[str] = []
        for source_path in sorted(archived_dir.iterdir(), key=lambda path: path.name.lower()):
            if not source_path.is_file():
                continue
            if hasattr(self.doc_manager, "is_supported_file") and not self.doc_manager.is_supported_file(source_path.name):
                skipped.append(source_path.name)
                continue
            safe_name = unique_input_filename(source_path.name, input_dir)
            target_path = input_dir / safe_name
            await asyncio.to_thread(shutil.copy2, source_path, target_path)
            copied_paths.append(target_path)

        if not copied_paths:
            await self.audit.record(
                principal=principal,
                action=ACTIVITY_DOCUMENT_UPLOAD,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                result="no_reindexable_files",
                metadata={"source": "__enqueued__", "skipped": skipped},
            )
            return LittleBullReindexArchivedResponse(
                status="no_files",
                message="Archived files were found, but none are supported for reindexing.",
                workspace_id=workspace_id,
                skipped_count=len(skipped),
            )

        track_id = generate_track_id("little_bull_reindex_archived")
        self._queue_pipeline_index_files(background_tasks, copied_paths, track_id, rag=rag)
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_UPLOAD,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="reindex_queued",
            metadata={
                "source": "__enqueued__",
                "track_id": track_id,
                "file_count": len(copied_paths),
                "skipped": skipped,
            },
        )
        return LittleBullReindexArchivedResponse(
            status="queued",
            message="Archived files were copied back and queued for reindexing.",
            track_id=track_id,
            workspace_id=workspace_id,
            recovered_count=len(copied_paths),
            skipped_count=len(skipped),
            files=[path.name for path in copied_paths],
        )

    async def query(self, principal: Principal, request: LittleBullQueryRequest) -> LittleBullQueryResponse:
        self._require(principal, ACTIVITY_QUERY, request.workspace_id)
        rag = await self._require_data_plane(request.workspace_id)
        tenant_id = await self._workspace_tenant(request.workspace_id)
        agent_config = await self._agent_config_for_query(
            tenant_id=tenant_id,
            workspace_id=request.workspace_id,
            agent_id=request.agent_id,
        )
        await self._require_agent_runtime_model_setting(
            tenant_id=tenant_id,
            workspace_id=request.workspace_id,
            agent_config=agent_config,
        )
        effective_model_profile = self._effective_model_profile(request.model_profile, agent_config)
        workspace_contains_private_data = bool(
            await self.repository.get_policy(
                WORKSPACE_PRIVATE_POLICY,
                tenant_id=tenant_id,
                workspace_id=request.workspace_id,
            )
        )
        hosted_private_policy = await self.repository.get_policy(
            PRIVATE_DATA_HOSTED_LLM_EXCEPTION_POLICY,
            tenant_id=tenant_id,
            workspace_id=request.workspace_id,
        )
        route_decision = self.private_gateway.evaluate(
            tenant_id=tenant_id,
            workspace_id=request.workspace_id,
            confidentiality=request.confidentiality,
            requested_profile=effective_model_profile,
            workspace_contains_private_data=workspace_contains_private_data,
            strict=private_strict_enabled(),
            hosted_private_policy=hosted_private_policy,
        )
        if not route_decision.allowed:
            await self.audit.record(
                principal=principal,
                action=ACTIVITY_QUERY,
                tenant_id=tenant_id,
                workspace_id=request.workspace_id,
                result="blocked",
                model=effective_model_profile,
                metadata={
                    **route_decision.audit_metadata(),
                    "confidentiality": request.confidentiality,
                    "workspace_contains_private_data": workspace_contains_private_data,
                },
            )
            raise HTTPException(
                status_code=route_decision.status_code or status.HTTP_403_FORBIDDEN,
                detail=route_decision.reason,
            )

        cache_disabled = route_decision.requires_private_runtime or route_decision.hosted_private_exception
        if cache_disabled:
            await self.audit.record(
                principal=principal,
                action=ACTIVITY_QUERY,
                tenant_id=tenant_id,
                workspace_id=request.workspace_id,
                result="allowed",
                model=effective_model_profile,
                metadata={
                    **route_decision.audit_metadata(),
                    "cache_disabled": cache_disabled,
                    "confidentiality": request.confidentiality,
                    "workspace_contains_private_data": workspace_contains_private_data,
                },
            )

        param = QueryParam(
            mode=self._effective_query_mode(request.mode, agent_config),
            response_type=self._effective_response_type(request.response_type, agent_config),
            stream=False,
            conversation_history=request.conversation_history,
        )
        if agent_config:
            param.user_prompt = build_agent_studio_prompt(agent_config)
        budget = await self._agent_context_budget_for_query(
            tenant_id=tenant_id,
            workspace_id=request.workspace_id,
            agent_config=agent_config,
        )
        budget_cost_state = await self._agent_context_budget_cost_state(
            tenant_id=tenant_id,
            workspace_id=request.workspace_id,
            agent_config=agent_config,
            budget=budget,
            query=request.query,
            user_prompt=getattr(param, "user_prompt", ""),
            conversation_history=request.conversation_history,
        )
        try:
            budget_metadata = self._enforce_agent_context_budget(
                budget=budget,
                query=request.query,
                user_prompt=getattr(param, "user_prompt", ""),
                conversation_history=request.conversation_history,
                estimated_request_cost_usd=budget_cost_state["estimated_request_cost_usd"],
                daily_cost_used_usd=budget_cost_state["daily_cost_used_usd"],
                monthly_cost_used_usd=budget_cost_state["monthly_cost_used_usd"],
            )
        except HTTPException as exc:
            await self.audit.record(
                principal=principal,
                action=ACTIVITY_QUERY,
                tenant_id=tenant_id,
                workspace_id=request.workspace_id,
                result="blocked",
                model=effective_model_profile,
                metadata={"reason": "agent_context_budget", "detail": exc.detail, "agent_id": request.agent_id},
            )
            raise
        self._apply_agent_context_budget_to_query_param(param, budget_metadata)
        if route_decision.requires_private_runtime:
            if route_decision.model_func is not None:
                param.model_func = route_decision.model_func
        elif not route_decision.hosted_private_exception:
            model_func = await self._model_func_for_profile(
                tenant_id=tenant_id,
                workspace_id=request.workspace_id,
                model_profile=effective_model_profile,
                agent_config=agent_config,
            )
            if model_func is not None:
                param.model_func = model_func
        if request.top_k is not None:
            param.top_k = request.top_k
        usage_ledger_id = await self._reserve_agent_query_budget_ledger(
            tenant_id=tenant_id,
            workspace_id=request.workspace_id,
            user_id=principal.user_id,
            agent_config=agent_config,
            budget=budget,
            budget_metadata=budget_metadata,
            effective_model_profile=effective_model_profile,
            query=request.query,
        )
        result = await self._aquery_with_private_cache_guard(
            request.query,
            param,
            private_runtime=cache_disabled,
            rag=rag,
        )
        llm_response = result.get("llm_response", {}) if isinstance(result, dict) else {}
        data = result.get("data", {}) if isinstance(result, dict) else {}
        response_content = llm_response.get("content") or "No relevant context found for the query."
        references = data.get("references", []) if request.include_references else []
        if request.include_references and request.include_chunk_content:
            references = self._enrich_references(references, data.get("chunks", []))
        if usage_ledger_id is None:
            usage_ledger_id = await self._append_agent_query_usage_ledger(
                tenant_id=tenant_id,
                workspace_id=request.workspace_id,
                user_id=principal.user_id,
                agent_config=agent_config,
                budget_metadata=budget_metadata,
                effective_model_profile=effective_model_profile,
                query=request.query,
                response=response_content,
            )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_QUERY,
            tenant_id=tenant_id,
            workspace_id=request.workspace_id,
            result="success",
            model=effective_model_profile,
            metadata={
                "mode": param.mode,
                "agent_id": request.agent_id,
                "usage_ledger_id": usage_ledger_id,
                "requested_model_profile": request.model_profile,
                "effective_model_profile": effective_model_profile,
                "reference_count": len(references),
                "agent_context_budget": budget_metadata,
                "private_gateway": {
                    **route_decision.audit_metadata(),
                    "cache_disabled": cache_disabled,
                },
            },
        )
        return LittleBullQueryResponse(
            response=response_content,
            references=references,
            workspace_id=request.workspace_id,
            model_profile=effective_model_profile,
        )

    async def _agent_context_budget_for_query(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        agent_config: dict[str, Any] | None,
    ) -> dict[str, Any] | None:
        if not agent_config or not agent_config.get("agent_id"):
            return None
        if not hasattr(self.admin_store, "list_agent_context_budgets"):
            return None
        budgets = await self.admin_store.list_agent_context_budgets(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            agent_id=agent_config["agent_id"],
        )
        if not budgets:
            return None
        model_setting_id = agent_config.get("model_setting_id")
        exact = next((budget for budget in budgets if budget.get("model_setting_id") == model_setting_id), None)
        default = next((budget for budget in budgets if budget.get("model_setting_id") is None), None)
        return exact or default or budgets[0]

    async def _reserve_agent_query_budget_ledger(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        user_id: str,
        agent_config: dict[str, Any] | None,
        budget: dict[str, Any] | None,
        budget_metadata: dict[str, Any],
        effective_model_profile: str,
        query: str,
    ) -> str | None:
        if not agent_config or not budget_metadata or not budget:
            return None
        daily_limit = self._float_or_none(budget.get("daily_cost_limit_usd"))
        monthly_limit = self._float_or_none(budget.get("monthly_cost_limit_usd"))
        if daily_limit is None and monthly_limit is None:
            return None
        if not hasattr(self.admin_store, "reserve_llm_usage_budget"):
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="LLM usage ledger is unavailable for atomic budget reservation.",
            )
        payload = await self._agent_query_usage_payload(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=user_id,
            agent_config=agent_config,
            budget_metadata=budget_metadata,
            effective_model_profile=effective_model_profile,
            query=query,
            response="",
            metadata_extra={"status": "reserved"},
        )
        today = utc_now().replace(hour=0, minute=0, second=0, microsecond=0)
        month_start = today.replace(day=1)
        try:
            row = await self.admin_store.reserve_llm_usage_budget(
                payload,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                user_id=user_id,
                agent_id=agent_config.get("agent_id"),
                daily_limit_usd=daily_limit,
                monthly_limit_usd=monthly_limit,
                daily_since=today,
                monthly_since=month_start,
            )
        except ValueError as exc:
            reason = str(exc)
            if "daily_cost_budget_exceeded" in reason:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Agent query exceeds daily cost budget.",
                ) from exc
            if "monthly_cost_budget_exceeded" in reason:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Agent query exceeds monthly cost budget.",
                ) from exc
            raise
        return row.get("usage_ledger_id")

    async def _append_agent_query_usage_ledger(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        user_id: str,
        agent_config: dict[str, Any] | None,
        budget_metadata: dict[str, Any],
        effective_model_profile: str,
        query: str,
        response: str,
    ) -> str | None:
        if not agent_config or not budget_metadata:
            return None
        if not hasattr(self.admin_store, "insert_llm_usage_ledger"):
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="LLM usage ledger is unavailable for budget debit.",
            )
        payload = await self._agent_query_usage_payload(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=user_id,
            agent_config=agent_config,
            budget_metadata=budget_metadata,
            effective_model_profile=effective_model_profile,
            query=query,
            response=response,
        )
        row = await self.admin_store.insert_llm_usage_ledger(
            payload,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=user_id,
        )
        return row.get("usage_ledger_id")

    async def _agent_query_usage_payload(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        user_id: str,
        agent_config: dict[str, Any],
        budget_metadata: dict[str, Any],
        effective_model_profile: str,
        query: str,
        response: str,
        metadata_extra: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        model_setting = await self._model_setting_for_agent_config(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            agent_config=agent_config,
        )
        prompt_tokens = int(budget_metadata.get("estimated_prompt_tokens") or 0) + int(
            budget_metadata.get("available_context_tokens") or 0
        )
        completion_tokens = int(budget_metadata.get("reserved_response_tokens") or 0)
        total_tokens = prompt_tokens + completion_tokens
        request_hash = "sha256:" + hashlib.sha256(query.encode("utf-8")).hexdigest()
        response_hash = "sha256:" + hashlib.sha256(response.encode("utf-8")).hexdigest()
        return {
            "user_id": user_id,
            "agent_id": agent_config.get("agent_id"),
            "model_setting_id": (model_setting or {}).get("model_setting_id"),
            "provider": (model_setting or {}).get("provider") or "runtime",
            "model_id": (model_setting or {}).get("model_id") or effective_model_profile,
            "operation": "agent_query",
            "prompt_tokens": prompt_tokens,
            "completion_tokens": completion_tokens,
            "total_tokens": total_tokens,
            "estimated_cost_usd": float(budget_metadata.get("estimated_request_cost_usd") or 0),
            "request_hash": request_hash,
            "response_hash": response_hash,
            "metadata": {
                "agent_context_budget_id": budget_metadata.get("agent_context_budget_id"),
                "effective_model_profile": effective_model_profile,
                **(metadata_extra or {}),
            },
        }

    async def _agent_context_budget_cost_state(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        agent_config: dict[str, Any] | None,
        budget: dict[str, Any] | None,
        query: str,
        user_prompt: str,
        conversation_history: list[dict[str, Any]],
    ) -> dict[str, float | None]:
        if not budget:
            return {
                "estimated_request_cost_usd": None,
                "daily_cost_used_usd": 0.0,
                "monthly_cost_used_usd": 0.0,
            }
        model_setting = await self._model_setting_for_agent_config(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            agent_config=agent_config,
        )
        prompt_tokens = self._estimate_agent_prompt_tokens(
            query=query,
            user_prompt=user_prompt,
            conversation_history=conversation_history,
        )
        reserved_response_tokens = int(budget.get("reserved_response_tokens") or 0)
        max_context_tokens = int(budget.get("max_context_tokens") or 0)
        available_context_tokens = max(0, max_context_tokens - prompt_tokens - reserved_response_tokens)
        estimated_cost = self._estimate_agent_request_cost_usd(
            budget=budget,
            model_setting=model_setting,
            input_tokens=prompt_tokens + available_context_tokens,
            output_tokens=reserved_response_tokens,
        )
        today = utc_now().replace(hour=0, minute=0, second=0, microsecond=0)
        month_start = today.replace(day=1)
        daily_used = await self._sum_agent_usage_cost(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            agent_id=(agent_config or {}).get("agent_id"),
            since=today,
        )
        monthly_used = await self._sum_agent_usage_cost(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            agent_id=(agent_config or {}).get("agent_id"),
            since=month_start,
        )
        return {
            "estimated_request_cost_usd": estimated_cost,
            "daily_cost_used_usd": daily_used,
            "monthly_cost_used_usd": monthly_used,
        }

    async def _model_setting_for_agent_config(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        agent_config: dict[str, Any] | None,
    ) -> dict[str, Any] | None:
        model_setting_id = (agent_config or {}).get("model_setting_id")
        if not model_setting_id:
            return None
        settings = await self._model_settings_for_workspace(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        return next((setting for setting in settings if setting.get("model_setting_id") == model_setting_id), None)

    async def _require_agent_runtime_model_setting(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        agent_config: dict[str, Any] | None,
    ) -> None:
        model_setting = await self._model_setting_for_agent_config(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            agent_config=agent_config,
        )
        model_setting_id = (agent_config or {}).get("model_setting_id")
        if not model_setting_id:
            return
        if not model_setting:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Agent runtime model setting not found.")
        if not model_setting.get("enabled", True) or model_setting.get("usage") not in {"chat", "agent"}:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Agent runtime model setting must be enabled and use chat or agent.",
            )

    async def _sum_agent_usage_cost(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        agent_id: str | None,
        since: datetime,
    ) -> float:
        if not hasattr(self.admin_store, "sum_llm_usage_cost"):
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="LLM usage ledger is unavailable for budget enforcement.",
            )
        return float(
            await self.admin_store.sum_llm_usage_cost(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                agent_id=agent_id,
                since=since,
            )
        )

    @staticmethod
    def _estimate_agent_prompt_tokens(
        *,
        query: str,
        user_prompt: str,
        conversation_history: list[dict[str, Any]],
    ) -> int:
        history_text = "\n".join(str(item.get("content") or "") for item in conversation_history)
        return estimate_tokens_from_characters(len("\n".join([query, user_prompt, history_text])))

    @classmethod
    def _estimate_agent_request_cost_usd(
        cls,
        *,
        budget: dict[str, Any],
        model_setting: dict[str, Any] | None,
        input_tokens: int,
        output_tokens: int,
    ) -> float | None:
        policy = budget.get("policy") or {}
        explicit_estimate = cls._float_or_none(policy.get("estimated_request_cost_usd"))
        if explicit_estimate is not None:
            return explicit_estimate
        pricing = dict((model_setting or {}).get("config") or {})
        pricing.update(policy)
        input_per_token = cls._float_or_none(pricing.get("input_price"))
        output_per_token = cls._float_or_none(pricing.get("output_price"))
        request_price = cls._float_or_none(pricing.get("request_price")) or 0.0
        if input_per_token is None:
            input_per_million = cls._float_or_none(
                pricing.get("prompt_cost_per_million_tokens")
                or pricing.get("input_cost_per_million_tokens")
                or pricing.get("cost_per_million_tokens")
            )
            input_per_token = input_per_million / 1_000_000 if input_per_million is not None else None
        if output_per_token is None:
            output_per_million = cls._float_or_none(
                pricing.get("completion_cost_per_million_tokens")
                or pricing.get("output_cost_per_million_tokens")
                or pricing.get("cost_per_million_tokens")
            )
            output_per_token = output_per_million / 1_000_000 if output_per_million is not None else None
        if input_per_token is None and output_per_token is None and not request_price:
            return None
        return round(
            (input_per_token or 0.0) * max(0, input_tokens)
            + (output_per_token or 0.0) * max(0, output_tokens)
            + request_price,
            8,
        )

    @staticmethod
    def _float_or_none(value: Any) -> float | None:
        if value is None or value == "":
            return None
        try:
            return float(value)
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _enforce_agent_context_budget(
        *,
        budget: dict[str, Any] | None,
        query: str,
        user_prompt: str,
        conversation_history: list[dict[str, Any]],
        estimated_request_cost_usd: float | None,
        daily_cost_used_usd: float,
        monthly_cost_used_usd: float,
    ) -> dict[str, Any]:
        if not budget:
            return {}
        prompt_tokens = LittleBullService._estimate_agent_prompt_tokens(
            query=query,
            user_prompt=user_prompt,
            conversation_history=conversation_history,
        )
        reserved_response_tokens = int(budget.get("reserved_response_tokens") or 0)
        max_prompt_tokens = int(budget.get("max_prompt_tokens") or 0)
        max_context_tokens = int(budget.get("max_context_tokens") or 0)
        available_context_tokens = 0
        if max_prompt_tokens and prompt_tokens > max_prompt_tokens:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Agent prompt exceeds max_prompt_tokens budget.",
            )
        if max_context_tokens and prompt_tokens + reserved_response_tokens > max_context_tokens:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Agent prompt plus reserved response exceeds max_context_tokens budget.",
            )
        if max_context_tokens:
            available_context_tokens = max_context_tokens - prompt_tokens - reserved_response_tokens
        daily_limit = LittleBullService._float_or_none(budget.get("daily_cost_limit_usd"))
        monthly_limit = LittleBullService._float_or_none(budget.get("monthly_cost_limit_usd"))
        if (daily_limit is not None or monthly_limit is not None) and not max_context_tokens:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Cost-limited agent context budgets require max_context_tokens.",
            )
        if (daily_limit is not None or monthly_limit is not None) and estimated_request_cost_usd is None:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Agent cost budget requires pricing metadata before runtime query.",
            )
        if daily_limit == 0 or monthly_limit == 0:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Agent cost budget is exhausted.",
            )
        if daily_limit is not None and daily_cost_used_usd + (estimated_request_cost_usd or 0.0) > daily_limit:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Agent query exceeds daily cost budget.",
            )
        if monthly_limit is not None and monthly_cost_used_usd + (estimated_request_cost_usd or 0.0) > monthly_limit:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Agent query exceeds monthly cost budget.",
            )
        return {
            "agent_context_budget_id": budget.get("agent_context_budget_id"),
            "estimated_prompt_tokens": prompt_tokens,
            "reserved_response_tokens": reserved_response_tokens,
            "max_prompt_tokens": max_prompt_tokens,
            "max_context_tokens": max_context_tokens,
            "available_context_tokens": available_context_tokens,
            "estimated_request_cost_usd": estimated_request_cost_usd,
            "daily_cost_used_usd": daily_cost_used_usd,
            "monthly_cost_used_usd": monthly_cost_used_usd,
        }

    @staticmethod
    def _apply_agent_context_budget_to_query_param(param: QueryParam, budget_metadata: dict[str, Any]) -> None:
        available_context_tokens = int(budget_metadata.get("available_context_tokens") or 0)
        if available_context_tokens <= 0:
            return
        param.max_total_tokens = min(int(getattr(param, "max_total_tokens", available_context_tokens)), available_context_tokens)
        param.max_entity_tokens = min(int(getattr(param, "max_entity_tokens", available_context_tokens)), available_context_tokens)
        param.max_relation_tokens = min(
            int(getattr(param, "max_relation_tokens", available_context_tokens)),
            available_context_tokens,
        )

    async def delete_document(self, principal: Principal, *, workspace_id: str, document_id: str) -> dict[str, Any]:
        self._require(principal, ACTIVITY_DOCUMENT_DELETE, workspace_id)
        rag = await self._require_data_plane(workspace_id)
        tenant_id = await self._workspace_tenant(workspace_id)
        if approvals_enforced():
            existing = await self._pending_delete_approval(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                document_id=document_id,
            )
            if existing is not None:
                await self.audit.record(
                    principal=principal,
                    action=ACTIVITY_DOCUMENT_DELETE,
                    tenant_id=tenant_id,
                    workspace_id=workspace_id,
                    result="pending_approval_existing",
                    approval_id=existing.approval_id,
                    metadata={"document_id": document_id},
                )
                return {
                    "status": "pending_approval",
                    "message": "Document deletion is already waiting for human approval.",
                    "approval": existing.to_dict(),
                }
            approval = await self.approvals.request(
                principal=principal,
                action=ACTIVITY_DOCUMENT_DELETE,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                reason="Document deletion requires human approval.",
                payload={"document_id": document_id},
            )
            await self.audit.record(
                principal=principal,
                action=ACTIVITY_DOCUMENT_DELETE,
                tenant_id=approval.tenant_id,
                workspace_id=workspace_id,
                result="pending_approval",
                approval_id=approval.approval_id,
                metadata={"document_id": document_id},
            )
            return {
                "status": "pending_approval",
                "message": "Document deletion is waiting for human approval.",
                "approval": approval.to_dict(),
            }
        if hasattr(rag, "adelete_by_doc_id"):
            await rag.adelete_by_doc_id(document_id)
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_DELETE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="success",
            metadata={"document_id": document_id},
        )
        return {"status": "success", "message": "Document deleted.", "doc_id": document_id}

    async def list_activity(self, principal: Principal, *, workspace_id: str, limit: int = 50) -> list[LittleBullActivityItem]:
        self._require(principal, ACTIVITY_ACTIVITY_READ, workspace_id)
        events = await self.audit.list(
            tenant_id=await self._workspace_tenant(workspace_id),
            workspace_id=workspace_id,
            limit=limit,
        )
        return [
            LittleBullActivityItem(
                id=event.event_id,
                action=event.action,
                result=event.result,
                created_at=event.created_at.isoformat(),
                actor_user_id=event.actor_user_id,
                workspace_id=event.workspace_id,
                metadata=event.metadata,
            )
            for event in events
        ]

    async def list_assistants(self, principal: Principal, *, workspace_id: str) -> list[LittleBullAssistant]:
        self._require(principal, ACTIVITY_ASSISTANTS_READ, workspace_id)
        tenant_id = await self._workspace_tenant(workspace_id)
        try:
            configs = await self.admin_store.list_agent_configs(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
        except Exception:
            configs = []
        if configs:
            return [
                LittleBullAssistant(
                    id=config["agent_id"],
                    name=config["name"],
                    description=config["description"],
                    enabled=config["enabled"],
                    response_rules=config["response_rules"],
                )
                for config in configs
            ]
        return [
            LittleBullAssistant(
                id="simple_answer",
                name="Resposta simples",
                description="Responde em linguagem direta usando o workspace ativo.",
                response_rules=["Usar fontes quando disponiveis", "Evitar resposta sem contexto"],
            ),
            LittleBullAssistant(
                id="checklist",
                name="Checklist",
                description="Transforma documentos em listas de providencias.",
                response_rules=["Mostrar itens acionaveis", "Preservar fontes"],
            ),
            LittleBullAssistant(
                id="private_local",
                name="Privado/local",
                description="Perfil exigido para dados sensiveis ou privados.",
                response_rules=["Nao usar modelo hospedado por padrao", "Registrar auditoria"],
            ),
        ]

    async def get_knowledge_graph(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        label: str,
        max_depth: int,
        max_nodes: int,
    ) -> Any:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        rag = await self._require_data_plane(workspace_id)
        tenant_id = await self._workspace_tenant(workspace_id)
        try:
            graph = await rag.get_knowledge_graph(
                node_label=label,
                max_depth=max_depth,
                max_nodes=max_nodes,
            )
        except Exception as exc:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(exc)) from exc
        await self.audit.record(
            principal=principal,
            action="little_bull.graph.read",
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="success",
            metadata={"label": label, "max_depth": max_depth, "max_nodes": max_nodes},
        )
        return graph

    async def list_graph_labels(self, principal: Principal, *, workspace_id: str) -> list[str]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        rag = await self._require_data_plane(workspace_id)
        try:
            return await rag.chunk_entity_relation_graph.get_all_labels()
        except Exception as exc:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(exc)) from exc

    async def list_popular_graph_labels(
        self, principal: Principal, *, workspace_id: str, limit: int
    ) -> list[str]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        rag = await self._require_data_plane(workspace_id)
        try:
            return await rag.chunk_entity_relation_graph.get_popular_labels(limit)
        except Exception as exc:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(exc)) from exc

    async def search_graph_labels(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        query: str,
        limit: int,
    ) -> list[str]:
        self._require(principal, ACTIVITY_DOCUMENT_READ, workspace_id)
        rag = await self._require_data_plane(workspace_id)
        try:
            return await rag.chunk_entity_relation_graph.search_labels(query, limit)
        except Exception as exc:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(exc)) from exc

    async def list_model_settings(
        self, principal: Principal, *, workspace_id: str
    ) -> list[LittleBullModelSetting]:
        self._require(principal, ACTIVITY_MODEL_MANAGE, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        settings = await self.admin_store.list_model_settings(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if not settings and await self._data_plane_attached(workspace_id):
            settings = self._runtime_model_defaults(workspace_id=workspace_id)
        return [LittleBullModelSetting(**setting) for setting in settings]

    async def list_embedding_catalog(self, principal: Principal) -> list[LittleBullEmbeddingCatalogItem]:
        self._require(principal, ACTIVITY_MODEL_MANAGE)
        self._require_master(principal)
        return [LittleBullEmbeddingCatalogItem(**item) for item in embedding_catalog()]

    async def list_knowledge_bases(self, principal: Principal) -> list[LittleBullKnowledgeBase]:
        self._require(principal, ACTIVITY_WORKSPACE_MANAGE)
        self._require_master(principal)
        workspaces = await self.repository.list_workspaces(None)
        bases: list[LittleBullKnowledgeBase] = []
        for workspace in workspaces:
            counts = await self._status_counts_for_workspace(workspace.workspace_id)
            settings = await self._model_settings_for_workspace(
                tenant_id=workspace.tenant_id,
                workspace_id=workspace.workspace_id,
            )
            chat_model = self._default_model(settings, "chat")
            embedding_model = self._default_model(settings, "embedding")
            estimated_tokens = await self._estimated_tokens_for_workspace(workspace.workspace_id)
            embedding_config = (embedding_model or {}).get("config", {})
            prompt_cost = float(embedding_config.get("prompt_cost_per_million_tokens") or 0)
            bases.append(
                LittleBullKnowledgeBase(
                    workspace_id=workspace.workspace_id,
                    tenant_id=workspace.tenant_id,
                    name=workspace.name,
                    slug=workspace.slug,
                    description=workspace.description,
                    privacy=workspace.privacy,
                    data_plane_attached=await self._data_plane_attached(workspace.workspace_id),
                    document_count=sum(counts.values()),
                    ready_count=counts.get("processed", 0),
                    processing_count=counts.get("processing", 0) + counts.get("pending", 0),
                    chat_model=LittleBullModelSetting(**chat_model) if chat_model else None,
                    embedding_model=LittleBullModelSetting(**embedding_model) if embedding_model else None,
                    embedding_reindex_required=bool(embedding_config.get("reindex_required")),
                    embedding_estimated_tokens=estimated_tokens,
                    embedding_estimated_cost_usd=estimate_embedding_cost(estimated_tokens, prompt_cost),
                )
            )
        return bases

    async def upsert_knowledge_base(
        self,
        principal: Principal,
        request: LittleBullKnowledgeBaseUpsertRequest,
    ) -> LittleBullKnowledgeBase:
        self._require(principal, ACTIVITY_WORKSPACE_MANAGE, request.workspace_id)
        self._require_master(principal)
        slug = slugify_workspace(request.slug or request.name)
        workspace_id = slugify_workspace(request.workspace_id or slug)
        tenant_id = principal.tenant_id or "default"
        workspace = Workspace(
            workspace_id=workspace_id,
            tenant_id=tenant_id,
            name=request.name.strip(),
            slug=slug,
            description=request.description.strip(),
            privacy=request.privacy.strip() or "team",
        )
        await self.repository.create_workspace(workspace)

        if request.embedding_model_id:
            await self._set_default_embedding_model(
                principal=principal,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                model_id=request.embedding_model_id,
                estimated_tokens=request.estimated_tokens,
            )

        await self.audit.record(
            principal=principal,
            action=ACTIVITY_WORKSPACE_MANAGE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="success",
            metadata={
                "slug": slug,
                "embedding_model_id": request.embedding_model_id,
                "data_plane_attached": await self._data_plane_attached(workspace_id),
            },
        )
        bases = await self.list_knowledge_bases(principal)
        return next(base for base in bases if base.workspace_id == workspace_id)

    async def attach_knowledge_base_data_plane(
        self,
        principal: Principal,
        *,
        workspace_id: str,
    ) -> LittleBullKnowledgeBaseAttachResponse:
        self._require(principal, ACTIVITY_WORKSPACE_MANAGE, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        rag = await self._ensure_workspace_data_plane(workspace_id)
        input_dir = self._input_dir_for_workspace(workspace_id)
        input_dir.mkdir(parents=True, exist_ok=True)
        policy = {
            "schema_version": 1,
            "attached": True,
            "workspace_id": workspace_id,
            "working_dir": str(getattr(rag, "working_dir", getattr(self.rag, "working_dir", "./rag_storage"))),
            "input_dir": str(input_dir),
            "attached_at": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
            "attached_by": principal.user_id,
        }
        await self.repository.set_policy(
            WORKSPACE_DATA_PLANE_POLICY,
            policy,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_WORKSPACE_MANAGE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="data_plane_attached",
            metadata=policy,
        )
        return LittleBullKnowledgeBaseAttachResponse(
            status="attached",
            message="Knowledge base is attached to a LightRAG data plane.",
            workspace_id=workspace_id,
            data_plane_attached=True,
            input_dir=str(input_dir),
            working_dir=policy["working_dir"],
        )

    async def reindex_knowledge_base(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        request: LittleBullKnowledgeBaseReindexRequest,
        background_tasks: BackgroundTasks,
    ) -> LittleBullKnowledgeBaseReindexResponse:
        self._require(principal, ACTIVITY_DOCUMENT_REINDEX, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        rag = await self._require_data_plane(workspace_id)
        approval = None
        payload = {
            "workspace_id": workspace_id,
            "include_archived": request.include_archived,
            "include_input_root": request.include_input_root,
            "destructive_rebuild": request.destructive_rebuild,
        }
        requires_approval = approvals_enforced() or request.destructive_rebuild

        if requires_approval:
            if not request.approval_id:
                existing = await self._pending_reindex_approval(
                    tenant_id=tenant_id,
                    workspace_id=workspace_id,
                    payload=payload,
                )
                if existing is not None:
                    await self.audit.record(
                        principal=principal,
                        action=ACTIVITY_DOCUMENT_REINDEX,
                        tenant_id=tenant_id,
                        workspace_id=workspace_id,
                        result="pending_approval_existing",
                        approval_id=existing.approval_id,
                        metadata=payload,
                    )
                    return LittleBullKnowledgeBaseReindexResponse(
                        status="pending_approval",
                        message="Knowledge base reindex is already waiting for human approval.",
                        workspace_id=workspace_id,
                        approval=existing.to_dict(),
                    )
                approval = await self.approvals.request(
                    principal=principal,
                    action=ACTIVITY_DOCUMENT_REINDEX,
                    tenant_id=tenant_id,
                    workspace_id=workspace_id,
                    reason=(
                        "Destructive knowledge base rebuild requires human approval."
                        if request.destructive_rebuild
                        else "Knowledge base reindex after embedding/model change requires human approval."
                    ),
                    payload=payload,
                )
                await self.audit.record(
                    principal=principal,
                    action=ACTIVITY_DOCUMENT_REINDEX,
                    tenant_id=tenant_id,
                    workspace_id=workspace_id,
                    result="pending_approval",
                    approval_id=approval.approval_id,
                    metadata=payload,
                )
                return LittleBullKnowledgeBaseReindexResponse(
                    status="pending_approval",
                    message="Knowledge base reindex is waiting for human approval.",
                    workspace_id=workspace_id,
                    approval=approval.to_dict(),
                )

            approval = await self.approvals.get(request.approval_id)
            if approval is None:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Approval not found.")
            if approval.action != ACTIVITY_DOCUMENT_REINDEX or approval.workspace_id != workspace_id:
                raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Approval does not match this reindex request.")
            approved_payload = approval.metadata or {}
            if (
                request.include_archived != coerce_payload_bool(approved_payload.get("include_archived"), True)
                or request.include_input_root != coerce_payload_bool(approved_payload.get("include_input_root"), True)
                or request.destructive_rebuild
                != coerce_payload_bool(approved_payload.get("destructive_rebuild"), False)
            ):
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail="Reindex request does not match the approved payload.",
                )
            if approval.status == ApprovalStatus.EXECUTED:
                return LittleBullKnowledgeBaseReindexResponse(
                    status="already_executed",
                    message="This approval has already been executed.",
                    workspace_id=workspace_id,
                    approval=approval.to_dict(),
                )
            if approval.status != ApprovalStatus.APPROVED:
                raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Approval must be approved before execution.")
            executing = await self.approvals.begin_execution(approval.approval_id, principal)
            approval = executing or await self.approvals.get(approval.approval_id) or approval

        snapshot_id = None
        snapshot_path = None
        drop_results: dict[str, Any] | None = None
        if request.destructive_rebuild:
            supported, skipped = await self._collect_reindex_sources(
                workspace_id=workspace_id,
                include_archived=request.include_archived,
                include_input_root=request.include_input_root,
            )
            if not supported:
                queued = LittleBullKnowledgeBaseReindexResponse(
                    status="no_files",
                    message=(
                        "No supported source files were found, so the destructive rebuild was not executed."
                    ),
                    workspace_id=workspace_id,
                    destructive_rebuild=True,
                    skipped_count=len(skipped),
                    files=[],
                )
            else:
                snapshot_id, snapshot_path = await self._snapshot_workspace_storage(
                    principal=principal,
                    tenant_id=tenant_id,
                    workspace_id=workspace_id,
                )
                drop_results = await self._drop_workspace_storages(rag)
                queued = self._queue_reindex_paths(
                    rag=rag,
                    workspace_id=workspace_id,
                    background_tasks=background_tasks,
                    paths=supported,
                    skipped=skipped,
                    destructive_rebuild=True,
                    snapshot_id=snapshot_id,
                    snapshot_path=snapshot_path,
                )
        else:
            queued = await self._queue_reindex_sources(
                rag=rag,
                workspace_id=workspace_id,
                background_tasks=background_tasks,
                include_archived=request.include_archived,
                include_input_root=request.include_input_root,
            )
        if approval is not None and approval.status == ApprovalStatus.EXECUTING:
            approval = await self.approvals.mark_executed(approval.approval_id, principal)
        await self._mark_embedding_reindex_queued(
            principal=principal,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            track_id=queued.track_id,
            queued_count=queued.queued_count,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_REINDEX,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result=queued.status,
            approval_id=approval.approval_id if approval else None,
            metadata={
                "track_id": queued.track_id,
                "queued_count": queued.queued_count,
                "skipped_count": queued.skipped_count,
                "files": queued.files,
                "destructive_rebuild": request.destructive_rebuild,
                "snapshot_id": snapshot_id,
                "snapshot_path": snapshot_path,
                "storage_drop_results": drop_results,
            },
        )
        if approval is not None:
            queued.approval = approval.to_dict()
        return queued

    async def rollback_knowledge_base_snapshot(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        request: LittleBullKnowledgeBaseRollbackRequest,
    ) -> LittleBullKnowledgeBaseRollbackResponse:
        self._require(principal, ACTIVITY_DOCUMENT_REINDEX, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        await self._require_data_plane(workspace_id)
        snapshot_id = self._safe_snapshot_id(request.snapshot_id)
        snapshot_dir = self._snapshot_root_for_workspace(workspace_id) / snapshot_id
        snapshot_storage_dir = snapshot_dir / "storage"
        if not snapshot_dir.exists() or not snapshot_storage_dir.exists():
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Snapshot not found.")
        if workspace_id == self._current_workspace_id():
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=(
                    "Online rollback of the active startup workspace is blocked. "
                    "Restore this snapshot while the server is stopped, then restart."
                ),
            )

        preserved_id, preserved_path = await self._snapshot_workspace_storage(
            principal=principal,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            reason="pre_rollback",
        )
        await self._restore_workspace_storage_from_snapshot(
            workspace_id=workspace_id,
            snapshot_storage_dir=snapshot_storage_dir,
        )
        self._workspace_rag_cache().pop(workspace_id, None)
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_DOCUMENT_REINDEX,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="rollback_restored",
            metadata={
                "snapshot_id": snapshot_id,
                "snapshot_path": str(snapshot_dir),
                "preserved_current_snapshot_id": preserved_id,
                "preserved_current_snapshot_path": preserved_path,
            },
        )
        return LittleBullKnowledgeBaseRollbackResponse(
            status="restored",
            message="Knowledge base snapshot restored. The workspace data plane will be recreated on next use.",
            workspace_id=workspace_id,
            snapshot_id=snapshot_id,
            restored_path=str(self._workspace_storage_dir(workspace_id)),
            preserved_current_snapshot_id=preserved_id,
            preserved_current_snapshot_path=preserved_path,
        )

    async def estimate_embedding_cost_for_workspace(
        self,
        principal: Principal,
        request: LittleBullEmbeddingCostEstimateRequest,
    ) -> LittleBullEmbeddingCostEstimateResponse:
        self._require(principal, ACTIVITY_MODEL_MANAGE, request.workspace_id)
        self._require_master(principal)
        entry = find_embedding_model(request.model_id)
        if entry is None:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Embedding model '{request.model_id}' is not in the OpenRouter catalog.",
            )
        notes: list[str] = []
        if request.estimated_tokens is not None:
            estimated_tokens = request.estimated_tokens
        elif request.page_count is not None:
            estimated_tokens = estimate_tokens_from_pages(
                request.page_count,
                request.words_per_page,
            )
            notes.append("Estimate based on page count and average words per page.")
        else:
            estimated_tokens = await self._estimated_tokens_for_workspace(request.workspace_id)
            notes.append("Estimate based on indexed document content length when the data plane is attached.")
        if not await self._data_plane_attached(request.workspace_id):
            notes.append("Workspace is configured but not attached to the current LightRAG data plane.")
        return LittleBullEmbeddingCostEstimateResponse(
            workspace_id=request.workspace_id,
            model_id=entry.model_id,
            display_name=entry.display_name,
            estimated_tokens=estimated_tokens,
            estimated_cost_usd=estimate_embedding_cost(
                estimated_tokens,
                entry.prompt_cost_per_million_tokens,
            ),
            prompt_cost_per_million_tokens=entry.prompt_cost_per_million_tokens,
            context_length=entry.context_length,
            recommended_chunk_tokens=entry.recommended_chunk_tokens,
            reindex_required=True,
            notes=notes,
        )

    async def _set_default_embedding_model(
        self,
        *,
        principal: Principal,
        tenant_id: str | None,
        workspace_id: str,
        model_id: str,
        estimated_tokens: int | None = None,
    ) -> dict[str, Any]:
        entry = find_embedding_model(model_id)
        if entry is None:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Embedding model '{model_id}' is not in the OpenRouter catalog.",
            )
        settings = await self._model_settings_for_workspace(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        current = self._default_model(settings, "embedding")
        current_config = dict((current or {}).get("config", {}) or {})
        changed = current is None or current.get("model_id") != entry.model_id
        tokens = estimated_tokens if estimated_tokens is not None else await self._estimated_tokens_for_workspace(workspace_id)
        config = {
            **current_config,
            "runtime_default": False,
            "requires_reindex": True,
            "reindex_required": bool(changed or current_config.get("reindex_required")),
            "prompt_cost_per_million_tokens": entry.prompt_cost_per_million_tokens,
            "context_length": entry.context_length,
            "recommended_chunk_tokens": entry.recommended_chunk_tokens,
            "estimated_reindex_tokens": tokens,
            "estimated_reindex_cost_usd": estimate_embedding_cost(tokens, entry.prompt_cost_per_million_tokens),
            "runtime_note": "Changing embeddings affects new indexing; existing vectors require reindexing.",
        }
        payload = {
            "model_setting_id": f"embedding_default_{workspace_id}",
            "tenant_id": tenant_id,
            "workspace_id": workspace_id,
            "usage": "embedding",
            "provider": "openrouter",
            "binding": entry.binding,
            "binding_host": entry.binding_host or OPENROUTER_EMBEDDING_HOST,
            "model_id": entry.model_id,
            "display_name": entry.display_name,
            "enabled": True,
            "is_default": True,
            "config": config,
        }
        return await self.admin_store.upsert_model_setting(
            payload,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )

    async def upsert_model_setting(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullModelSetting,
    ) -> LittleBullModelSetting:
        self._require(principal, ACTIVITY_MODEL_MANAGE, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        if payload.model_setting_id:
            settings = await self._model_settings_for_workspace(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            existing = next(
                (setting for setting in settings if setting.get("model_setting_id") == payload.model_setting_id),
                None,
            )
            if not existing:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Model setting not found.")
            if existing.get("tenant_id") != tenant_id or existing.get("workspace_id") != workspace_id:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Model settings cannot be moved across workspace scope.",
                )
        if payload.usage == "embedding" and payload.is_default:
            settings = await self._model_settings_for_workspace(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            current = self._default_model(settings, "embedding")
            payload.config["reindex_required"] = bool(
                current is None
                or current.get("model_id") != payload.model_id
                or payload.config.get("reindex_required")
            )
            payload.config["requires_reindex"] = True
            payload.config.setdefault(
                "runtime_note",
                "Changing embeddings affects new indexing; existing vectors require reindexing.",
            )
        try:
            saved = await self.admin_store.upsert_model_setting(
                payload.model_dump(exclude_none=True),
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                user_id=principal.user_id,
            )
        except ValueError as exc:
            if "model_setting_scope_mismatch" not in str(exc):
                raise
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Model settings cannot be moved across workspace scope.",
            ) from exc
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_MODEL_MANAGE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="success",
            model=saved.get("model_id"),
            metadata={"usage": saved.get("usage"), "provider": saved.get("provider")},
        )
        return LittleBullModelSetting(**saved)

    async def list_agent_configs(
        self, principal: Principal, *, workspace_id: str
    ) -> list[LittleBullAgentConfig]:
        self._require(principal, ACTIVITY_AGENT_MANAGE, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        configs = await self.admin_store.list_agent_configs(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if not configs:
            configs = self._default_agent_configs(workspace_id=workspace_id)
        return [LittleBullAgentConfig(**self._normalize_agent_config(config)) for config in configs]

    async def upsert_agent_config(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullAgentConfig,
    ) -> LittleBullAgentConfig:
        self._require(principal, ACTIVITY_AGENT_MANAGE, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        if payload.agent_id:
            existing_agents = await self.admin_store.list_agent_configs(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not any(agent.get("agent_id") == payload.agent_id for agent in existing_agents):
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Agent not found in workspace.")
        payload_data = self._normalize_agent_config(payload.model_dump(exclude_none=True))
        await self._require_scoped_model_setting(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            model_setting_id=payload_data.get("model_setting_id"),
            usage={"chat", "agent"},
        )
        issues, score = validate_agent_studio_config(payload_data)
        errors = [issue for issue in issues if issue["severity"] == "error"]
        if errors:
            await self.audit.record(
                principal=principal,
                action=ACTIVITY_AGENT_MANAGE,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                result="blocked",
                metadata={
                    "agent_id": payload.agent_id,
                    "readiness_score": score,
                    "errors": errors,
                },
            )
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail={
                    "message": "Agent Studio validation blocked publish.",
                    "readiness_score": score,
                    "issues": issues,
                },
            )
        try:
            saved = await self.admin_store.upsert_agent_config(
                payload_data,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                user_id=principal.user_id,
            )
        except ValueError as exc:
            if "agent_config_scope_mismatch" not in str(exc):
                raise
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Existing agents cannot be moved across workspace scope.",
            ) from exc
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_AGENT_MANAGE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="success",
            metadata={"agent_id": saved.get("agent_id"), "model_setting_id": saved.get("model_setting_id")},
        )
        return LittleBullAgentConfig(**self._normalize_agent_config(saved))

    async def _require_scoped_model_setting(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        model_setting_id: str | None,
        usage: str | set[str] | None,
    ) -> None:
        if not model_setting_id:
            return
        settings = await self._model_settings_for_workspace(tenant_id=tenant_id, workspace_id=workspace_id)
        allowed_usage = {usage} if isinstance(usage, str) else usage
        for setting in settings:
            if setting.get("model_setting_id") != model_setting_id:
                continue
            if allowed_usage and setting.get("usage") not in allowed_usage:
                raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Model usage mismatch.")
            return
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Model setting not found in workspace.")

    async def preview_agent_studio(
        self,
        principal: Principal,
        request: LittleBullAgentStudioPreviewRequest,
    ) -> LittleBullAgentStudioPreviewResponse:
        self._require(principal, ACTIVITY_AGENT_MANAGE, request.workspace_id)
        self._require_master(principal)
        await self._existing_workspace_scope(request.workspace_id)
        preview = agent_studio_preview(request.agent.model_dump(exclude_none=True))
        test_summary = "Validação estática concluída."
        if request.test_input.strip():
            test_summary = (
                "Teste rápido compilou o prompt e verificou governança básica; "
                "nenhuma chamada de modelo foi executada nesta validação."
            )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_AGENT_MANAGE,
            tenant_id=await self._workspace_tenant(request.workspace_id),
            workspace_id=request.workspace_id,
            result="preview",
            metadata={
                "agent_id": request.agent.agent_id,
                "readiness_score": preview["readiness_score"],
                "issue_count": len(preview["issues"]),
            },
        )
        return LittleBullAgentStudioPreviewResponse(
            agent=LittleBullAgentConfig(**preview["agent"]),
            issues=preview["issues"],
            readiness_score=preview["readiness_score"],
            ready_to_publish=preview["ready_to_publish"],
            compiled_prompt=preview["compiled_prompt"],
            test_input=request.test_input,
            test_summary=test_summary,
        )

    async def list_agent_builder_sessions(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        status_filter: str | None = None,
    ) -> list[AgentBuilderSession]:
        self._require(principal, ACTIVITY_AGENT_MANAGE, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        rows = await self.admin_store.list_agent_builder_sessions(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=None,
            status=status_filter,
        )
        return [AgentBuilderSession(**row) for row in rows]

    async def upsert_agent_builder_session(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullAgentBuilderSessionRequest,
    ) -> AgentBuilderSession:
        self._require(principal, ACTIVITY_AGENT_MANAGE, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        await self._require_scoped_model_setting(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            model_setting_id=payload.model_setting_id,
            usage="agent_builder",
        )
        existing = None
        transcript: list[dict[str, Any]] = []
        if payload.agent_builder_session_id:
            existing = await self.admin_store.get_agent_builder_session(
                payload.agent_builder_session_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            if not existing:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Agent builder session not found.")
            transcript = list(existing.get("builder_transcript") or [])
        transcript.append({"role": "user", "content": payload.user_message.strip(), "created_at": utc_now()})
        generated_config = self._agent_builder_generated_config(
            workspace_id=workspace_id,
            user_message=payload.user_message,
            existing_config=payload.generated_config or (existing or {}).get("generated_config") or {},
        )
        await self._require_scoped_model_setting(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            model_setting_id=generated_config.get("model_setting_id"),
            usage={"chat", "agent"},
        )
        preview = agent_studio_preview(generated_config)
        transcript.append(
            {
                "role": "assistant",
                "content": "Agent Builder draft updated for human review.",
                "readiness_score": preview["readiness_score"],
                "created_at": utc_now(),
            }
        )
        row = await self.admin_store.upsert_agent_builder_session(
            {
                "agent_builder_session_id": payload.agent_builder_session_id,
                "user_id": principal.user_id,
                "agent_id": (existing or {}).get("agent_id"),
                "model_setting_id": payload.model_setting_id or (existing or {}).get("model_setting_id"),
                "status": "draft",
                "current_step": payload.current_step,
                "builder_transcript": transcript,
                "generated_config": preview["agent"],
                "readiness_score": preview["readiness_score"],
                "requires_review": True,
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_AGENT_MANAGE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="agent_builder_draft",
            metadata={
                "agent_builder_session_id": row["agent_builder_session_id"],
                "readiness_score": row["readiness_score"],
            },
        )
        return AgentBuilderSession(**row)

    async def publish_agent_builder_session(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        agent_builder_session_id: str,
        payload: LittleBullAgentBuilderPublishRequest,
    ) -> AgentBuilderSession:
        self._require(principal, ACTIVITY_AGENT_MANAGE, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        session = await self.admin_store.get_agent_builder_session(
            agent_builder_session_id,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
        )
        if not session:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Agent builder session not found.")
        if not payload.approved:
            raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Human approval is required to publish.")
        generated_config = self._normalize_agent_config(session.get("generated_config") or {})
        generated_config["enabled"] = payload.enabled
        await self._require_scoped_model_setting(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            model_setting_id=generated_config.get("model_setting_id"),
            usage={"chat", "agent"},
        )
        preview = agent_studio_preview(generated_config)
        if not preview["ready_to_publish"]:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail={
                    "message": "Agent Builder draft is not ready to publish.",
                    "readiness_score": preview["readiness_score"],
                    "issues": preview["issues"],
                },
            )
        saved = await self.admin_store.upsert_agent_config(
            preview["agent"],
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        transcript = list(session.get("builder_transcript") or [])
        transcript.append(
            {
                "role": "system",
                "content": "Human-approved draft published.",
                "agent_id": saved["agent_id"],
                "created_at": utc_now(),
            }
        )
        row = await self.admin_store.upsert_agent_builder_session(
            {
                **session,
                "agent_id": saved["agent_id"],
                "status": "published",
                "builder_transcript": transcript,
                "generated_config": preview["agent"],
                "readiness_score": preview["readiness_score"],
                "requires_review": False,
            },
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_AGENT_MANAGE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="agent_builder_published",
            metadata={"agent_builder_session_id": agent_builder_session_id, "agent_id": saved["agent_id"]},
        )
        return AgentBuilderSession(**row)

    @staticmethod
    def _agent_builder_generated_config(
        *,
        workspace_id: str,
        user_message: str,
        existing_config: dict[str, Any],
    ) -> dict[str, Any]:
        title = str(existing_config.get("name") or user_message.strip().splitlines()[0]).strip()[:80]
        if not title:
            title = "Agente Little Bull"
        config = normalize_agent_studio_config(existing_config.get("config"))
        config["identity"].update(
            {
                "mission": config["identity"].get("mission") or user_message.strip()[:500],
                "when_to_use": config["identity"].get("when_to_use") or "Quando a tarefa estiver dentro do escopo configurado.",
                "when_not_to_use": config["identity"].get("when_not_to_use") or "Quando faltar contexto ou aprovação humana.",
                "audience": config["identity"].get("audience") or "Usuários do workspace Little Bull.",
            }
        )
        config["knowledge"]["allowed_workspace_ids"] = [workspace_id]
        config["tests"] = config.get("tests") or [
            {
                "name": "Recusa sem contexto",
                "input": "Responda sem fontes.",
                "expected_behavior": "Solicitar contexto e indicar incerteza.",
                "forbidden_behavior": "Inventar fontes.",
            }
        ]
        return {
            "name": title,
            "description": str(existing_config.get("description") or user_message.strip()[:240]).strip(),
            "enabled": False,
            "model_setting_id": existing_config.get("model_setting_id"),
            "system_prompt": str(existing_config.get("system_prompt") or "").strip(),
            "response_rules": existing_config.get("response_rules") or ["Citar fontes quando usar conhecimento recuperado."],
            "tools": existing_config.get("tools") or ["query_knowledge"],
            "config": config,
        }

    async def list_agent_context_budgets(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        agent_id: str | None = None,
    ) -> list[AgentContextBudget]:
        self._require(principal, ACTIVITY_AGENT_MANAGE, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        rows = await self.admin_store.list_agent_context_budgets(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            agent_id=agent_id,
        )
        return [AgentContextBudget(**row) for row in rows]

    async def upsert_agent_context_budget(
        self,
        principal: Principal,
        *,
        workspace_id: str,
        payload: LittleBullAgentContextBudgetRequest,
    ) -> AgentContextBudget:
        self._require(principal, ACTIVITY_AGENT_MANAGE, workspace_id)
        self._require_master(principal)
        tenant_id, workspace_id = await self._existing_workspace_scope(workspace_id)
        agents = await self.admin_store.list_agent_configs(tenant_id=tenant_id, workspace_id=workspace_id)
        if not any(agent.get("agent_id") == payload.agent_id for agent in agents):
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Agent not found.")
        if payload.agent_context_budget_id:
            existing_budgets = await self.admin_store.list_agent_context_budgets(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            existing_budget = next(
                (
                    budget
                    for budget in existing_budgets
                    if budget["agent_context_budget_id"] == payload.agent_context_budget_id
                ),
                None,
            )
            if not existing_budget:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Agent context budget not found.")
            if existing_budget["agent_id"] != payload.agent_id:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Existing agent context budgets cannot be moved across agents.",
                )
        await self._require_scoped_model_setting(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            model_setting_id=payload.model_setting_id,
            usage=None,
        )
        if payload.reserved_response_tokens > payload.max_context_tokens:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="reserved_response_tokens cannot exceed max_context_tokens.",
            )
        if payload.max_prompt_tokens and payload.max_prompt_tokens > (
            payload.max_context_tokens - payload.reserved_response_tokens
        ):
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="max_prompt_tokens must fit inside context after reserved response tokens.",
            )
        if (payload.daily_cost_limit_usd is not None or payload.monthly_cost_limit_usd is not None) and (
            payload.max_context_tokens <= 0
        ):
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Cost-limited agent context budgets require max_context_tokens.",
            )
        try:
            row = await self.admin_store.upsert_agent_context_budget(
                payload.model_dump(exclude_none=True),
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                user_id=principal.user_id,
            )
        except ValueError as exc:
            if "agent_context_budget_scope_mismatch" not in str(exc):
                raise
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Existing agent context budgets cannot be moved across agent/workspace scope.",
            ) from exc
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_AGENT_MANAGE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="agent_context_budget_upserted",
            metadata={"agent_id": row["agent_id"], "agent_context_budget_id": row["agent_context_budget_id"]},
        )
        return AgentContextBudget(**row)

    async def save_conversation(
        self,
        principal: Principal,
        request: LittleBullConversationSaveRequest,
    ) -> LittleBullConversation:
        self._require(principal, ACTIVITY_CONVERSATION_SAVE, request.workspace_id)
        tenant_id, workspace_id = await self._scope_for_workspace(request.workspace_id)
        saved = await self.admin_store.save_conversation(
            request.model_dump(exclude_none=True),
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_CONVERSATION_SAVE,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="success",
            metadata={
                "conversation_id": saved["conversation_id"],
                "message_count": saved.get("message_count", 0),
            },
        )
        return LittleBullConversation(**saved)

    async def list_conversations(
        self, principal: Principal, *, workspace_id: str
    ) -> list[LittleBullConversation]:
        self._require(principal, ACTIVITY_CONVERSATION_READ, workspace_id)
        tenant_id, workspace_id = await self._scope_for_workspace(workspace_id)
        user_id = None if principal.is_master_global else principal.user_id
        conversations = await self.admin_store.list_conversations(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=user_id,
        )
        return [LittleBullConversation(**conversation) for conversation in conversations]

    async def get_conversation(
        self, principal: Principal, *, conversation_id: str
    ) -> LittleBullConversation:
        conversation = await self.admin_store.get_conversation(conversation_id)
        if not conversation:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Conversation not found")
        self._require(principal, ACTIVITY_CONVERSATION_READ, conversation["workspace_id"])
        if not principal.is_master_global and conversation["user_id"] != principal.user_id:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Conversation belongs to another user.")
        return LittleBullConversation(**conversation)

    async def export_conversation(
        self,
        principal: Principal,
        *,
        conversation_id: str,
        export_format: str,
    ) -> Response:
        conversation = await self.get_conversation(principal, conversation_id=conversation_id)
        self._require(principal, ACTIVITY_CONVERSATION_EXPORT, conversation.workspace_id)
        normalized = export_format.lower().strip()
        if normalized == "md":
            body = self._conversation_markdown(conversation)
            media_type = "text/markdown; charset=utf-8"
            suffix = "md"
            content = body.encode("utf-8")
        elif normalized == "txt":
            body = self._conversation_text(conversation)
            media_type = "text/plain; charset=utf-8"
            suffix = "txt"
            content = body.encode("utf-8")
        elif normalized == "docx":
            content = self._conversation_docx(conversation)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            suffix = "docx"
        else:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported export format")
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_CONVERSATION_EXPORT,
            tenant_id=conversation.tenant_id,
            workspace_id=conversation.workspace_id,
            result="success",
            metadata={"conversation_id": conversation_id, "format": normalized},
        )
        filename = f"little-bull-{conversation.conversation_id}.{suffix}"
        return Response(
            content=content,
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    async def create_correlation_suggestion(
        self,
        principal: Principal,
        request: LittleBullCorrelationSuggestionRequest,
    ) -> LittleBullCorrelationSuggestion:
        self._require(principal, ACTIVITY_CORRELATION_SUGGEST, request.workspace_id)
        tenant_id, workspace_id = await self._scope_for_workspace(request.workspace_id)
        if request.source_label.strip() == request.target_label.strip():
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Source and target must differ")
        saved = await self.admin_store.create_correlation_suggestion(
            request.model_dump(exclude_none=True),
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            user_id=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_CORRELATION_SUGGEST,
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            result="pending",
            metadata={
                "suggestion_id": saved["suggestion_id"],
                "source_label": saved["source_label"],
                "target_label": saved["target_label"],
            },
        )
        return LittleBullCorrelationSuggestion(**saved)

    async def list_correlation_suggestions(
        self, principal: Principal, *, workspace_id: str, suggestion_status: str | None = None
    ) -> list[LittleBullCorrelationSuggestion]:
        self._require(principal, ACTIVITY_CORRELATION_SUGGEST, workspace_id)
        tenant_id, workspace_id = await self._scope_for_workspace(workspace_id)
        suggestions = await self.admin_store.list_correlation_suggestions(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            status=suggestion_status,
        )
        return [LittleBullCorrelationSuggestion(**suggestion) for suggestion in suggestions]

    async def decide_correlation_suggestion(
        self,
        principal: Principal,
        *,
        suggestion_id: str,
        decision: str,
    ) -> LittleBullCorrelationSuggestion:
        self._require(principal, ACTIVITY_CORRELATION_DECIDE)
        normalized = decision.strip().lower()
        if normalized not in {"approved", "rejected"}:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid decision")
        current = await self.admin_store.get_correlation_suggestion(suggestion_id)
        if not current:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Suggestion not found")
        self._require(principal, ACTIVITY_CORRELATION_DECIDE, current["workspace_id"])
        saved = await self.admin_store.decide_correlation_suggestion(
            suggestion_id,
            status=normalized,
            decided_by=principal.user_id,
        )
        await self.audit.record(
            principal=principal,
            action=ACTIVITY_CORRELATION_DECIDE,
            tenant_id=saved["tenant_id"],
            workspace_id=saved["workspace_id"],
            result=normalized,
            metadata={"suggestion_id": suggestion_id},
        )
        return LittleBullCorrelationSuggestion(**saved)

    async def _pending_delete_approval(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        document_id: str,
    ) -> Any | None:
        pending = await self.approvals.list(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            status=ApprovalStatus.PENDING,
        )
        for approval in pending:
            if approval.action != ACTIVITY_DOCUMENT_DELETE:
                continue
            approved_document = approval.metadata.get("document_id") or approval.metadata.get("doc_id")
            if str(approved_document or "") == document_id:
                return approval
        return None

    async def _pending_reindex_approval(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        payload: dict[str, Any],
    ) -> Any | None:
        pending = await self.approvals.list(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            status=ApprovalStatus.PENDING,
        )
        for approval in pending:
            if approval.action != ACTIVITY_DOCUMENT_REINDEX:
                continue
            metadata = approval.metadata or {}
            if (
                str(metadata.get("workspace_id") or "") == str(payload.get("workspace_id") or "")
                and coerce_payload_bool(metadata.get("include_archived"), True)
                == coerce_payload_bool(payload.get("include_archived"), True)
                and coerce_payload_bool(metadata.get("include_input_root"), True)
                == coerce_payload_bool(payload.get("include_input_root"), True)
                and coerce_payload_bool(metadata.get("destructive_rebuild"), False)
                == coerce_payload_bool(payload.get("destructive_rebuild"), False)
            ):
                return approval
        return None

    async def _queue_reindex_sources(
        self,
        *,
        rag: Any,
        workspace_id: str,
        background_tasks: BackgroundTasks,
        include_archived: bool,
        include_input_root: bool,
    ) -> LittleBullKnowledgeBaseReindexResponse:
        supported, skipped = await self._collect_reindex_sources(
            workspace_id=workspace_id,
            include_archived=include_archived,
            include_input_root=include_input_root,
        )
        return self._queue_reindex_paths(
            rag=rag,
            workspace_id=workspace_id,
            background_tasks=background_tasks,
            paths=supported,
            skipped=skipped,
            destructive_rebuild=False,
        )

    async def _collect_reindex_sources(
        self,
        *,
        workspace_id: str,
        include_archived: bool,
        include_input_root: bool,
    ) -> tuple[list[Path], list[str]]:
        input_dir = self._input_dir_for_workspace(workspace_id)
        input_dir.mkdir(parents=True, exist_ok=True)
        candidates: list[Path] = []
        skipped: list[str] = []
        if include_input_root:
            for path in sorted(input_dir.iterdir(), key=lambda item: item.name.lower()):
                if path.is_file():
                    candidates.append(path)
        if include_archived:
            archived_dir = input_dir / "__enqueued__"
            if archived_dir.exists():
                for source_path in sorted(archived_dir.iterdir(), key=lambda item: item.name.lower()):
                    if not source_path.is_file():
                        continue
                    if hasattr(self.doc_manager, "is_supported_file") and not self.doc_manager.is_supported_file(source_path.name):
                        skipped.append(source_path.name)
                        continue
                    target_path = input_dir / unique_input_filename(source_path.name, input_dir)
                    await asyncio.to_thread(shutil.copy2, source_path, target_path)
                    candidates.append(target_path)

        supported: list[Path] = []
        seen: set[Path] = set()
        for path in candidates:
            if path in seen:
                continue
            seen.add(path)
            if hasattr(self.doc_manager, "is_supported_file") and not self.doc_manager.is_supported_file(path.name):
                skipped.append(path.name)
                continue
            supported.append(path)

        return supported, skipped

    def _queue_reindex_paths(
        self,
        *,
        rag: Any,
        workspace_id: str,
        background_tasks: BackgroundTasks,
        paths: list[Path],
        skipped: list[str],
        destructive_rebuild: bool,
        snapshot_id: str | None = None,
        snapshot_path: str | None = None,
    ) -> LittleBullKnowledgeBaseReindexResponse:
        supported = paths
        if not supported:
            return LittleBullKnowledgeBaseReindexResponse(
                status="no_files",
                message="No supported source files were found to queue for reindexing.",
                workspace_id=workspace_id,
                destructive_rebuild=destructive_rebuild,
                snapshot_id=snapshot_id,
                snapshot_path=snapshot_path,
                rollback_available=bool(snapshot_id),
                skipped_count=len(skipped),
            )

        track_id = generate_track_id("little_bull_reindex_base")
        self._queue_pipeline_index_files(background_tasks, supported, track_id, rag=rag)
        message = (
            "Snapshot was created, existing index data was dropped, and source files were queued for rebuild."
            if destructive_rebuild
            else (
                "Source files were queued for reindexing. Existing duplicate content may still require "
                "a destructive rebuild before vectors are replaced."
            )
        )
        return LittleBullKnowledgeBaseReindexResponse(
            status="queued",
            message=message,
            workspace_id=workspace_id,
            track_id=track_id,
            destructive_rebuild=destructive_rebuild,
            snapshot_id=snapshot_id,
            snapshot_path=snapshot_path,
            rollback_available=bool(snapshot_id),
            queued_count=len(supported),
            skipped_count=len(skipped),
            files=[path.name for path in supported],
        )

    def _workspace_storage_dir(self, workspace_id: str) -> Path:
        working_root = Path(getattr(self.rag, "working_dir", "./rag_storage")).expanduser()
        raw_current_workspace = str(getattr(self.rag, "workspace", "") or "")
        if workspace_id == self._current_workspace_id() and not raw_current_workspace:
            candidate = working_root
        else:
            candidate = working_root / workspace_id
        root_resolved = working_root.resolve()
        candidate_resolved = candidate.resolve()
        if candidate_resolved != root_resolved and not candidate_resolved.is_relative_to(root_resolved):
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Unsafe workspace storage path.",
            )
        return candidate

    def _snapshot_root_for_workspace(self, workspace_id: str) -> Path:
        safe_workspace = slugify_workspace(workspace_id)
        return Path(getattr(self.rag, "working_dir", "./rag_storage")).expanduser() / "__little_bull_snapshots__" / safe_workspace

    @staticmethod
    def _safe_snapshot_id(snapshot_id: str) -> str:
        value = snapshot_id.strip()
        if not re.fullmatch(r"[A-Za-z0-9_.-]+", value):
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid snapshot id.")
        return value

    async def _snapshot_workspace_storage(
        self,
        *,
        principal: Principal,
        tenant_id: str | None,
        workspace_id: str,
        reason: str = "pre_rebuild",
    ) -> tuple[str, str]:
        storage_dir = self._workspace_storage_dir(workspace_id)
        snapshot_id = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S%fZ")
        snapshot_dir = self._snapshot_root_for_workspace(workspace_id) / snapshot_id
        snapshot_storage_dir = snapshot_dir / "storage"
        await asyncio.to_thread(snapshot_dir.mkdir, parents=True, exist_ok=False)
        if storage_dir.exists():
            await asyncio.to_thread(
                shutil.copytree,
                storage_dir,
                snapshot_storage_dir,
                ignore=shutil.ignore_patterns("__little_bull_snapshots__"),
            )
        else:
            await asyncio.to_thread(snapshot_storage_dir.mkdir, parents=True, exist_ok=True)
        metadata = {
            "schema_version": 1,
            "snapshot_id": snapshot_id,
            "workspace_id": workspace_id,
            "tenant_id": tenant_id,
            "reason": reason,
            "storage_dir": str(storage_dir),
            "created_by": principal.user_id,
            "created_at": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        }
        await asyncio.to_thread(
            (snapshot_dir / "metadata.json").write_text,
            json.dumps(metadata, indent=2, sort_keys=True),
            encoding="utf-8",
        )
        return snapshot_id, str(snapshot_dir)

    async def _drop_workspace_storages(self, rag: Any) -> dict[str, Any]:
        results: dict[str, Any] = {}
        storage_names = [
            "full_docs",
            "text_chunks",
            "full_entities",
            "full_relations",
            "entity_chunks",
            "relation_chunks",
            "entities_vdb",
            "relationships_vdb",
            "chunks_vdb",
            "chunk_entity_relation_graph",
            "llm_response_cache",
            "doc_status",
        ]
        for storage_name in storage_names:
            storage_obj = getattr(rag, storage_name, None)
            drop = getattr(storage_obj, "drop", None)
            if not callable(drop):
                continue
            result = drop()
            if hasattr(result, "__await__"):
                result = await result
            results[storage_name] = result
            if isinstance(result, dict) and result.get("status") == "error":
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Could not drop workspace storage '{storage_name}': {result.get('message')}",
                )
        if not results:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="No droppable LightRAG storages are available for destructive rebuild.",
            )
        return results

    async def _restore_workspace_storage_from_snapshot(
        self,
        *,
        workspace_id: str,
        snapshot_storage_dir: Path,
    ) -> None:
        storage_dir = self._workspace_storage_dir(workspace_id)

        def restore() -> None:
            storage_dir.mkdir(parents=True, exist_ok=True)
            for child in storage_dir.iterdir():
                if child.name == "__little_bull_snapshots__":
                    continue
                if child.is_dir():
                    shutil.rmtree(child)
                else:
                    child.unlink()
            for child in snapshot_storage_dir.iterdir():
                target = storage_dir / child.name
                if child.is_dir():
                    shutil.copytree(child, target)
                else:
                    shutil.copy2(child, target)

        await asyncio.to_thread(restore)

    async def _mark_embedding_reindex_queued(
        self,
        *,
        principal: Principal,
        tenant_id: str | None,
        workspace_id: str,
        track_id: str | None,
        queued_count: int,
    ) -> None:
        try:
            settings = await self._model_settings_for_workspace(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
            embedding_model = self._default_model(settings, "embedding")
            if not embedding_model:
                return
            config = dict(embedding_model.get("config", {}) or {})
            config["last_reindex_queued_at"] = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
            config["last_reindex_track_id"] = track_id
            config["last_reindex_file_count"] = queued_count
            embedding_model["config"] = config
            await self.admin_store.upsert_model_setting(
                embedding_model,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                user_id=principal.user_id,
            )
        except Exception:
            return

    async def _documents_paginated(self, *, rag: Any, page: int, page_size: int) -> tuple[tuple[list[tuple[str, Any]], int], dict[str, int]]:
        docs_task = rag.doc_status.get_docs_paginated(
            status_filter=None,
            page=page,
            page_size=page_size,
            sort_field="updated_at",
            sort_direction="desc",
        )
        counts_task = rag.doc_status.get_all_status_counts()
        documents = await docs_task
        counts = await counts_task
        return documents, dict(counts or {})

    def _queue_pipeline_index_files(
        self,
        background_tasks: BackgroundTasks,
        copied_paths: list[Path],
        track_id: str,
        *,
        rag: Any,
    ) -> None:
        from lightrag.api.routers.document_routes import pipeline_index_files

        background_tasks.add_task(pipeline_index_files, rag, copied_paths, track_id)

    def _queue_pipeline_index_file(
        self,
        background_tasks: BackgroundTasks,
        file_path: Path,
        track_id: str,
        *,
        rag: Any,
    ) -> None:
        from lightrag.api.routers.document_routes import pipeline_index_file

        background_tasks.add_task(pipeline_index_file, rag, file_path, track_id)

    async def _agent_config_for_query(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        agent_id: str | None,
    ) -> dict[str, Any] | None:
        if not agent_id:
            return None
        try:
            agents = await self.admin_store.list_agent_configs(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Agent store is unavailable.",
            ) from exc
        if not agents:
            agents = self._default_agent_configs(workspace_id=workspace_id)
        agent = next((agent for agent in agents if agent.get("agent_id") == agent_id and agent.get("enabled")), None)
        if not agent:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Agent not found or disabled.")
        return self._normalize_agent_config(agent)

    @staticmethod
    def _effective_model_profile(
        requested_profile: str,
        agent_config: dict[str, Any] | None,
    ) -> str:
        if not agent_config:
            return requested_profile
        model_config = normalize_agent_studio_config(
            agent_config.get("config"),
            agent_config.get("tools") or [],
        ).get("model", {})
        return str(model_config.get("profile") or requested_profile or "equilibrado").strip()

    @staticmethod
    def _effective_query_mode(
        requested_mode: str,
        agent_config: dict[str, Any] | None,
    ) -> str:
        if not agent_config or requested_mode != "mix":
            return requested_mode
        knowledge_config = normalize_agent_studio_config(
            agent_config.get("config"),
            agent_config.get("tools") or [],
        ).get("knowledge", {})
        return str(knowledge_config.get("retrieval_mode") or requested_mode).strip()

    @staticmethod
    def _effective_response_type(
        requested_response_type: str,
        agent_config: dict[str, Any] | None,
    ) -> str:
        if not agent_config or requested_response_type != "Multiple Paragraphs":
            return requested_response_type
        output_config = normalize_agent_studio_config(
            agent_config.get("config"),
            agent_config.get("tools") or [],
        ).get("output", {})
        format_to_response_type = {
            "texto": "Multiple Paragraphs",
            "paragrafos": "Multiple Paragraphs",
            "checklist": "Bullet Points",
            "bullets": "Bullet Points",
            "resumo": "Single Paragraph",
            "markdown": "Multiple Paragraphs",
        }
        output_format = str(output_config.get("default_format") or "").strip().lower()
        return format_to_response_type.get(output_format, requested_response_type)

    async def _model_func_for_profile(
        self,
        *,
        tenant_id: str | None,
        workspace_id: str,
        model_profile: str,
        agent_config: dict[str, Any] | None,
    ) -> Any | None:
        try:
            models = await self.admin_store.list_model_settings(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
            )
        except Exception:
            return None
        selected_id = agent_config.get("model_setting_id") if agent_config else None
        selected = None
        if selected_id:
            selected = next((model for model in models if model.get("model_setting_id") == selected_id), None)
            if selected is None:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Agent runtime model setting not found.")
            if not selected.get("enabled", True) or selected.get("usage") not in {"chat", "agent"}:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Agent runtime model setting must be enabled and use chat or agent.",
                )
        if selected is None:
            selected = next(
                (
                    model
                    for model in models
                    if model.get("usage") in {"chat", "agent"}
                    and model.get("enabled")
                    and (
                        model.get("model_setting_id") == model_profile
                        or model.get("config", {}).get("profile") == model_profile
                    )
                ),
                None,
            )
        if selected is None:
            selected = next(
                (
                    model
                    for model in models
                    if model.get("usage") in {"chat", "agent"} and model.get("enabled") and model.get("is_default")
                ),
                None,
            )
        if not selected:
            return None
        if selected.get("binding") != "openai":
            return None
        api_key_ref = str(selected.get("config", {}).get("api_key_ref") or "OPENROUTER_API_KEY")
        api_key_ref = api_key_ref.removeprefix("env:")
        api_key = os.getenv(api_key_ref) or os.getenv("LLM_BINDING_API_KEY") or os.getenv("OPENAI_API_KEY")
        if not api_key:
            return None
        agent_model_config = (
            normalize_agent_studio_config(agent_config.get("config"), agent_config.get("tools") or []).get("model", {})
            if agent_config
            else {}
        )
        model_kwargs: dict[str, Any] = {}
        if agent_model_config.get("temperature") is not None:
            model_kwargs["temperature"] = float(agent_model_config["temperature"])
        if agent_model_config.get("max_tokens"):
            model_kwargs["max_tokens"] = int(agent_model_config["max_tokens"])
        return partial(
            openai_complete_if_cache,
            selected["model_id"],
            base_url=selected.get("binding_host") or os.getenv("LLM_BINDING_HOST"),
            api_key=api_key,
            **model_kwargs,
        )

    def _runtime_model_defaults(self, *, workspace_id: str) -> list[dict[str, Any]]:
        llm_model = str(getattr(self.rag, "little_bull_llm_model", None) or os.getenv("LLM_MODEL") or "openai/gpt-4o-mini")
        llm_host = str(getattr(self.rag, "little_bull_llm_host", None) or os.getenv("LLM_BINDING_HOST") or "https://openrouter.ai/api/v1")
        embedding_model = str(os.getenv("EMBEDDING_MODEL") or getattr(self.rag, "embedding_model_name", None) or "runtime-embedding")
        embedding_binding = str(os.getenv("EMBEDDING_BINDING") or "openai")
        embedding_host = str(os.getenv("EMBEDDING_BINDING_HOST") or os.getenv("LLM_BINDING_HOST") or "")
        return [
            {
                "model_setting_id": "runtime_chat_default",
                "tenant_id": None,
                "workspace_id": workspace_id,
                "usage": "chat",
                "provider": "openrouter" if "openrouter.ai" in llm_host else "runtime",
                "binding": str(getattr(self.rag, "little_bull_llm_binding", None) or os.getenv("LLM_BINDING") or "openai"),
                "binding_host": llm_host,
                "model_id": llm_model,
                "display_name": f"Chat padrão ({llm_model})",
                "enabled": True,
                "is_default": True,
                "config": {"profile": "equilibrado", "api_key_ref": "env:OPENROUTER_API_KEY", "runtime_default": True},
            },
            {
                "model_setting_id": "runtime_embedding_default",
                "tenant_id": None,
                "workspace_id": workspace_id,
                "usage": "embedding",
                "provider": "openrouter" if "openrouter.ai" in embedding_host else "runtime",
                "binding": embedding_binding,
                "binding_host": embedding_host,
                "model_id": embedding_model,
                "display_name": f"Embedding padrão ({embedding_model})",
                "enabled": True,
                "is_default": True,
                "config": {
                    "runtime_default": True,
                    "requires_reindex": True,
                    "runtime_note": "Embeddings novos exigem reindexação para alterar vetores existentes.",
                },
            },
        ]

    @staticmethod
    def _default_agent_configs(*, workspace_id: str) -> list[dict[str, Any]]:
        configs = [
            {
                "agent_id": "simple_answer",
                "workspace_id": workspace_id,
                "name": "Resposta simples",
                "description": "Responde em linguagem direta usando o workspace ativo.",
                "enabled": True,
                "model_setting_id": None,
                "system_prompt": "Responda de forma direta, com fontes quando disponíveis, sem inventar contexto.",
                "response_rules": ["Usar fontes quando disponiveis", "Evitar resposta sem contexto"],
                "tools": ["query_knowledge"],
                "config": {"profile": "equilibrado"},
            },
            {
                "agent_id": "checklist",
                "workspace_id": workspace_id,
                "name": "Checklist",
                "description": "Transforma documentos em listas de providencias.",
                "enabled": True,
                "model_setting_id": None,
                "system_prompt": "Transforme a resposta em checklist acionável e preserve fontes relevantes.",
                "response_rules": ["Mostrar itens acionaveis", "Preservar fontes"],
                "tools": ["query_knowledge"],
                "config": {"profile": "equilibrado"},
            },
            {
                "agent_id": "private_local",
                "workspace_id": workspace_id,
                "name": "Privado/local",
                "description": "Perfil exigido para dados sensiveis ou privados.",
                "enabled": True,
                "model_setting_id": None,
                "system_prompt": "Trate dados privados com minimização e não use modelo hospedado sem política MASTER válida.",
                "response_rules": ["Nao usar modelo hospedado por padrao", "Registrar auditoria"],
                "tools": ["query_knowledge"],
                "config": {"profile": "privado"},
            },
        ]
        return [LittleBullService._normalize_agent_config(config) for config in configs]

    @staticmethod
    def _normalize_agent_config(agent: dict[str, Any]) -> dict[str, Any]:
        copy = dict(agent)
        copy["tools"] = [normalize_tool_id(item) for item in copy.get("tools") or [] if normalize_tool_id(item)]
        copy["response_rules"] = [
            str(item) for item in copy.get("response_rules") or [] if str(item).strip()
        ]
        copy["config"] = normalize_agent_studio_config(copy.get("config"), copy["tools"])
        return copy

    @staticmethod
    def _conversation_markdown(conversation: LittleBullConversation) -> str:
        lines = [f"# {conversation.title}", "", f"- Workspace: {conversation.workspace_id}", f"- Modelo: {conversation.model_profile}", ""]
        for message in conversation.messages:
            role = "Usuário" if message.role == "user" else "Assistente"
            lines.extend([f"## {role}", "", message.content, ""])
            for index, reference in enumerate(message.references, start=1):
                lines.append(f"- Fonte {index}: {reference.get('file_path') or reference.get('reference_id') or 'sem identificador'}")
            if message.references:
                lines.append("")
        return "\n".join(lines).strip() + "\n"

    @staticmethod
    def _conversation_text(conversation: LittleBullConversation) -> str:
        parts = [conversation.title, f"Workspace: {conversation.workspace_id}", f"Modelo: {conversation.model_profile}", ""]
        for message in conversation.messages:
            role = "Usuario" if message.role == "user" else "Assistente"
            parts.append(f"{role}: {message.content}")
            if message.references:
                refs = ", ".join(str(ref.get("file_path") or ref.get("reference_id") or "sem identificador") for ref in message.references)
                parts.append(f"Fontes: {refs}")
            parts.append("")
        return "\n".join(parts).strip() + "\n"

    @staticmethod
    def _conversation_docx(conversation: LittleBullConversation) -> bytes:
        try:
            from docx import Document
        except ImportError as exc:
            raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="python-docx is not installed") from exc

        document = Document()
        document.add_heading(conversation.title, level=1)
        document.add_paragraph(f"Workspace: {conversation.workspace_id}")
        document.add_paragraph(f"Modelo: {conversation.model_profile}")
        for message in conversation.messages:
            role = "Usuário" if message.role == "user" else "Assistente"
            document.add_heading(role, level=2)
            document.add_paragraph(message.content)
            for index, reference in enumerate(message.references, start=1):
                document.add_paragraph(
                    f"Fonte {index}: {reference.get('file_path') or reference.get('reference_id') or 'sem identificador'}",
                    style="List Bullet",
                )
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    async def _status_counts_safe(self, *, rag: Any) -> dict[str, int]:
        try:
            return dict(await rag.doc_status.get_all_status_counts())
        except Exception:
            return {}

    async def _aquery_with_private_cache_guard(
        self,
        query: str,
        param: QueryParam,
        *,
        private_runtime: bool,
        rag: Any,
    ) -> Any:
        if not private_runtime:
            return await rag.aquery_llm(query, param=param)
        cache = getattr(rag, "llm_response_cache", None)
        global_config = getattr(cache, "global_config", None)
        if not isinstance(global_config, dict) or "enable_llm_cache" not in global_config:
            return await rag.aquery_llm(query, param=param)
        previous = global_config.get("enable_llm_cache")
        global_config["enable_llm_cache"] = False
        try:
            return await rag.aquery_llm(query, param=param)
        finally:
            global_config["enable_llm_cache"] = previous

    @staticmethod
    def _enrich_references(references: list[dict[str, Any]], chunks: list[dict[str, Any]]) -> list[dict[str, Any]]:
        ref_id_to_content: dict[str, list[str]] = {}
        for chunk in chunks:
            ref_id = chunk.get("reference_id")
            content = chunk.get("content")
            if ref_id and content:
                ref_id_to_content.setdefault(ref_id, []).append(content)
        enriched: list[dict[str, Any]] = []
        for reference in references:
            copy = dict(reference)
            if copy.get("reference_id") in ref_id_to_content:
                copy["content"] = ref_id_to_content[copy["reference_id"]]
            enriched.append(copy)
        return enriched
